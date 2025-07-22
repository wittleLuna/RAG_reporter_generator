from fastapi import FastAPI, File, UploadFile, Form, HTTPException, Request, Body, Depends, status
from fastapi.responses import HTMLResponse, FileResponse, RedirectResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from fastapi.security import HTTPBasic, HTTPBasicCredentials
from starlette.middleware.sessions import SessionMiddleware
import uvicorn
import os
import uuid
import shutil
from pathlib import Path
import asyncio
from typing import List, Optional, Dict
import aiofiles
from dotenv import load_dotenv
import json
from datetime import datetime
import logging
from docx import Document
from docxtpl import DocxTemplate, RichText, InlineImage
from docx.shared import Mm, RGBColor, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
import re
from docxcompose.composer import Composer
import traceback
import secrets
from passlib.context import CryptContext
from sqlalchemy import create_engine, Column, Integer, String, Text, DateTime, Boolean, ForeignKey, or_
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import sessionmaker, Session
from sqlalchemy.sql import func
import hashlib
from sqlalchemy.orm import relationship
from openai import OpenAI
import base64
from backend.models import Message

# 加载环境变量
load_dotenv()

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('logs/app.log', encoding='utf-8'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# 数据库配置（提前）
DATABASE_URL = "sqlite:///./rag_system.db"
engine = create_engine(DATABASE_URL, connect_args={"check_same_thread": False})
SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)
Base = declarative_base()

def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

# 导入千问API版本的服务
from services.ai_service_latest import AIService

app = FastAPI(title="RAG实训报告生成系统", version="1.0.0")

# 添加会话中间件
app.add_middleware(SessionMiddleware, secret_key="your-secret-key-here")

# 挂载uploads静态目录
app.mount("/uploads", StaticFiles(directory="uploads"), name="uploads")

# 创建必要的目录
os.makedirs("uploads", exist_ok=True)
os.makedirs("temp", exist_ok=True)
os.makedirs("static", exist_ok=True)
os.makedirs("templates", exist_ok=True)
os.makedirs("chroma_db", exist_ok=True)
os.makedirs("logs", exist_ok=True)
os.makedirs("user_templates", exist_ok=True)  # 用户模板目录

# 聊天图片目录
os.makedirs("chat_images", exist_ok=True)

# 静态文件和模板配置
app.mount("/static", StaticFiles(directory="static"), name="static")
templates = Jinja2Templates(directory="templates")

# 初始化AI服务
ai_service = AIService()

# 定义支持的文本文件类型
TEXT_FILE_EXTENSIONS = {'.txt', '.md', '.markdown', '.py', '.js', '.html', '.css', '.json', '.xml', '.csv', '.log', '.ini', '.conf', '.yaml', '.yml'}

def is_text_file(file_path):
    """判断是否为文本文件"""
    return file_path.suffix.lower() in TEXT_FILE_EXTENSIONS

def read_text_file_content(file_path):
    """安全读取文本文件内容"""
    try:
        # 尝试多种编码
        encodings = ['utf-8', 'gbk', 'gb2312', 'utf-8-sig']
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                    # 检查内容是否包含大量不可打印字符（可能是二进制文件）
                    if len(content) > 0 and len([c for c in content if ord(c) < 32 and c not in '\n\r\t']) / len(content) < 0.1:
                        return content
            except UnicodeDecodeError:
                continue
        return None
    except Exception as e:
        logger.warning(f"无法读取文件 {file_path}: {e}")
        return None

def read_word_document_content(file_path):
    """读取Word文档内容"""
    try:
        doc = Document(file_path)
        content = []
        
        # 读取段落内容
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                content.append(paragraph.text.strip())
        
        # 读取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_content = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_content.append(cell.text.strip())
                if row_content:
                    content.append(" | ".join(row_content))
        
        return "\n".join(content)
    except Exception as e:
        logger.warning(f"无法读取Word文档 {file_path}: {e}")
        return None

def find_placeholder_paragraph(doc, placeholder_text="{{report_body}}"):
    """在文档中查找占位符段落，包括表格中的占位符"""
    # 1. 首先在段落中查找
    for paragraph in doc.paragraphs:
        if placeholder_text in paragraph.text:
            logger.info(f"在段落中找到占位符: {paragraph.text[:50]}...")
            return paragraph
    
    # 2. 在表格中查找
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if placeholder_text in paragraph.text:
                        logger.info(f"在表格单元格中找到占位符: {paragraph.text[:50]}...")
                        return paragraph
    
    # 3. 在页眉页脚中查找（如果需要）
    for section in doc.sections:
        # 页眉
        if section.header:
            for paragraph in section.header.paragraphs:
                if placeholder_text in paragraph.text:
                    logger.info(f"在页眉中找到占位符: {paragraph.text[:50]}...")
                    return paragraph
        
        # 页脚
        if section.footer:
            for paragraph in section.footer.paragraphs:
                if placeholder_text in paragraph.text:
                    logger.info(f"在页脚中找到占位符: {paragraph.text[:50]}...")
                    return paragraph
    
    logger.warning(f"未找到占位符: {placeholder_text}")
    return None

def get_style(doc, style_names):
    """获取存在的样式名，支持多个备选样式名"""
    for style_name in style_names:
        try:
            doc.styles[style_name]
            return style_name
        except KeyError:
            continue
    return None

def add_structured_content(doc, target_paragraph, markdown_text, image_dir="uploads", uploaded_images=None):
    """在指定段落位置添加结构化内容"""
    lines = markdown_text.split('\n')
    in_code_block = False
    code_lines = []
    
    # 获取可用的样式
    code_style = get_style(doc, ['Code'])
    heading1_style = get_style(doc, ['Heading 1', '标题 1'])
    heading2_style = get_style(doc, ['Heading 2', '标题 2'])
    heading3_style = get_style(doc, ['Heading 3', '标题 3'])
    heading4_style = get_style(doc, ['Heading 4', '标题 4'])
    list_style = get_style(doc, ['项目符号'])
    normal_style = get_style(doc, ['Normal', '正文'])
    bold_style = get_style(doc, ['bold', 'Bold', 'Strong'])
    
    # 删除原占位符段落
    if target_paragraph:
        p = target_paragraph._element
        p.getparent().remove(p)
    
    for line in lines:
        # 代码块识别（markdown风格）
        if line.strip().startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_lines = []
            else:
                # 代码块结束，渲染代码
                in_code_block = False
                code_text = '\n'.join(code_lines)
                # 创建代码块段落，优先使用自定义Code样式
                p = doc.add_paragraph()
                if code_style:
                    p.style = code_style
                    # 添加文本内容到已设置样式的段落
                    p.add_run(code_text)
                    logger.info(f"使用Code样式添加代码块: {len(code_lines)}行")
                else:
                    # 如果没有Code样式，手动设置字体
                    run = p.add_run(code_text)
                    run.font.name = "Courier New"
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(64, 64, 64)  # 灰色
                    logger.info(f"手动设置代码块样式: {len(code_lines)}行")
                code_lines.clear()
            continue
            
        if in_code_block:
            code_lines.append(line)
            continue
            
        # 行内图片占位符替换
        img_inline = re.findall(r'{{image:(img_\d+)}}', line)
        if img_inline and uploaded_images:
            p = doc.add_paragraph()
            p.style = normal_style or 'Normal'
            last_idx = 0
            for match in re.finditer(r'{{image:(img_\d+)}}', line):
                start, end = match.span()
                if start > last_idx:
                    p.add_run(line[last_idx:start])
                img_id = match.group(1)
                img_path = find_image_by_id(img_id, uploaded_images, image_dir)
                logger.info(f"尝试插入图片: {img_id} 路径: {img_path} 存在: {os.path.exists(img_path) if img_path else False}")
                if img_path and os.path.exists(img_path):
                    run = p.add_run()
                    run.add_picture(img_path, width=Inches(4))
                    logger.info(f"成功插入图片: {img_id} -> {img_path}")
                else:
                    p.add_run(f"[图片{img_id}未找到]")
                last_idx = end
            if last_idx < len(line):
                p.add_run(line[last_idx:])
            continue
            
        # 项目符号
        if line.strip().startswith('- '):
            p = doc.add_paragraph(line.strip()[2:])
            p.style = list_style or 'Normal'
        # 四级标题
        elif line.strip().startswith('#### '):
            p = doc.add_paragraph(line.strip().replace('#### ', ''))
            p.style = heading4_style or 'Normal'
        # 三级标题
        elif line.strip().startswith('### '):
            p = doc.add_paragraph(line.strip().replace('### ', ''))
            p.style = heading3_style or 'Normal'
        # 二级标题
        elif line.strip().startswith('## '):
            p = doc.add_paragraph(line.strip().replace('## ', ''))
            p.style = heading2_style or 'Normal'
        # 一级标题
        elif line.strip().startswith('# '):
            p = doc.add_paragraph(line.strip().replace('# ', ''))
            p.style = heading1_style or 'Normal'
        # 普通段落（包含粗体处理）
        elif line.strip():
            # 处理粗体文本
            if '**' in line:
                # 分割文本，处理粗体部分
                parts = line.split('**')
                p = doc.add_paragraph()
                p.style = normal_style or 'Normal'
                
                for i, part in enumerate(parts):
                    if part.strip():  # 跳过空字符串
                        if i % 2 == 1:  # 奇数索引是粗体文本
                            run = p.add_run(part)
                            # 优先使用字符样式，如果没有则使用字体属性
                            if bold_style and bold_style != 'Normal':
                                try:
                                    # 检查是否为字符样式
                                    if doc.styles[bold_style].type == WD_STYLE_TYPE.CHARACTER:
                                        run.style = bold_style
                                        logger.info(f"使用字符样式应用粗体: {part}")
                                    else:
                                        # 如果不是字符样式，使用字体属性
                                        run.bold = True
                                        logger.info(f"使用字体属性应用粗体: {part}")
                                except Exception as e:
                                    # 如果样式应用失败，使用字体属性
                                    run.bold = True
                                    logger.warning(f"样式应用失败，使用字体属性: {e}")
                            else:
                                # 没有找到粗体样式，使用字体属性
                                run.bold = True
                                logger.info(f"使用字体属性应用粗体: {part}")
                        else:  # 偶数索引是普通文本
                            p.add_run(part)
            else:
                p = doc.add_paragraph(line.strip())
                p.style = normal_style or 'Normal'

def find_image_by_id(img_id, uploaded_images, image_dir="uploads"):
    for img in uploaded_images:
        if img['id'] == img_id:
            # 修正：拼接uploads/用户ID/图片名
            return os.path.join(image_dir, img['filepath'])
    return None

def find_image_path(img_name, upload_dir):
    """查找图片文件路径"""
    norm_img_name = normalize_filename(img_name)
    for fname in os.listdir(upload_dir):
        if normalize_filename(fname) == norm_img_name:
            return os.path.join(upload_dir, fname)
    return None

def normalize_filename(name):
    """标准化文件名用于匹配"""
    return os.path.splitext(name.strip().lower().replace(' ', '').replace('_', ''))[0]

@app.get("/", response_class=HTMLResponse)
async def index(request: Request, db: Session = Depends(get_db)):
    """主页"""
    user = get_current_user(request, db)
    if not user:
        return RedirectResponse(url="/login", status_code=303)
    user_dict = {
        "id": user.id,
        "username": user.username,
        "email": user.email,
        "name": user.name,
        "student_id": user.student_id,
        "class_name": user.class_name
    }
    return templates.TemplateResponse("index.html", {"request": request, "user": user_dict})

@app.get("/test-images", response_class=HTMLResponse)
async def test_images(request: Request):
    """图片渲染测试页面"""
    return templates.TemplateResponse("test_images.html", {"request": request})

@app.get("/health")
async def health():
    """健康检查"""
    return {"status": "healthy", "timestamp": datetime.now().isoformat()}

@app.post("/upload")
async def upload_files(
    request: Request,
    cover_template: UploadFile = File(None),
    body_template: UploadFile = File(None),
    files: list[UploadFile] = File([]),
    db: Session = Depends(get_db)
):
    # 获取当前用户
    user = get_current_user(request, db)
    if not user:
        raise HTTPException(status_code=401, detail="请先登录")
    user_dir = f"uploads/{user.id}"
    os.makedirs(user_dir, exist_ok=True)
    report_image_dir = f"report_images/{user.id}"
    os.makedirs(report_image_dir, exist_ok=True)
    uploaded_files = []
    images = []
    if cover_template and cover_template.filename:
        file_path = f"{user_dir}/cover_template_{cover_template.filename}"
        async with aiofiles.open(file_path, 'wb') as f:
            content = await cover_template.read()
            await f.write(content)
        uploaded_files.append({"filename": cover_template.filename, "type": "cover_template", "path": file_path})
    if body_template and body_template.filename:
        file_path = f"{user_dir}/body_template_{body_template.filename}"
        async with aiofiles.open(file_path, 'wb') as f:
            content = await body_template.read()
            await f.write(content)
        uploaded_files.append({"filename": body_template.filename, "type": "body_template", "path": file_path})
    for file in files:
        if file.filename:
            file_path = f"{user_dir}/{file.filename}"
            async with aiofiles.open(file_path, 'wb') as f:
                content = await file.read()
                await f.write(content)
            uploaded_files.append({"filename": file.filename, "type": "data_file", "path": file_path})
            # 如果是图片，复制到 report_images/用户ID/
            if file.filename.lower().endswith((".jpg", ".jpeg", ".png", ".gif", ".bmp")):
                import shutil
                dst_path = f"{report_image_dir}/{file.filename}"
                shutil.copy(file_path, dst_path)
                images.append({
                    "filename": file.filename,
                    "webpath": f"report_images/{user.id}/{file.filename}",
                    "path": dst_path
                })
    return {"success": True, "files": uploaded_files, "images": images}

def count_pages(docx_path):
    """估算Word文档页数（字符数+分页符）"""
    try:
        doc = Document(docx_path)
        total_chars = 0
        for paragraph in doc.paragraphs:
            total_chars += len(paragraph.text)
        estimated_pages = max(1, total_chars // 600)
        page_breaks = 0
        for paragraph in doc.paragraphs:
            if paragraph.runs:
                for run in paragraph.runs:
                    if hasattr(run, '_element') and run._element.xml.find('w:br') != -1:
                        page_breaks += 1
        return max(estimated_pages, page_breaks + 1)
    except Exception as e:
        logger.warning(f"页数估算失败: {e}")
        return 0

async def generate_report_to_target_pages(query, context_text, target_pages, max_rounds=8):
    """多轮补全+自动扩写，直到接近目标页数，优化版"""
    all_content = ""
    current_pages = 0
    for round_num in range(max_rounds):
        if round_num == 0:
            prompt = build_prompt(query, context_text, target_pages=target_pages)
            round_tokens = 8000 if target_pages and target_pages > 6 else 4000
        else:
            # 续写时强制要求补充字数且不重复
            prompt = (
                f"请在以下内容基础上，继续补充详细内容，要求总长度达到{target_pages}页Word文档，"
                f"本轮请补充不少于{min(2000, (target_pages-current_pages)*600)}字的新内容，"
                "增加更多技术细节、案例、分析、数据、图表等，内容要丰富、详实，且不要重复前文。\n\n"
                f"【已生成内容】\n{all_content}"
            )
            round_tokens = 8000 if target_pages and target_pages > 6 else 4000
        logger.info(f"\n==== 多轮补全第{round_num+1}轮 ====")
        logger.info(f"Prompt片段: {prompt[:200]}...")
        new_content = await ai_service.generate_report_with_prompt(query, prompt, target_pages)
        logger.info(f"本轮AI返回内容长度: {len(new_content)} 字")
        if not new_content.strip():
            logger.warning("AI本轮未返回内容，提前终止")
            break
        if new_content.strip() in all_content:
            logger.warning("AI输出内容重复，提前终止补全")
            break
        all_content += "\n" + new_content.strip()
        # 写入临时docx统计页数
        temp_docx = "temp/_temp_check.docx"
        from docx import Document
        doc = Document()
        for para in all_content.split('\n'):
            doc.add_paragraph(para)
        doc.save(temp_docx)
        current_pages = count_pages(temp_docx)
        logger.info(f"当前总字数: {len(all_content)}，当前页数: {current_pages}")
        logger.info(f"本轮补全文本片段: {new_content[:200]}...")
        if current_pages >= target_pages * 0.9:
            logger.info(f"已达到目标页数({current_pages}/{target_pages})，终止补全")
            break
    logger.info(f"==== 多轮补全结束，最终总字数: {len(all_content)}，最终页数: {current_pages} ====")
    return all_content.strip()

@app.post("/generate_report")
async def generate_report(
    request: Request,
    query: str = Form(...),
    name: str = Form(None),
    student_id: str = Form(None),
    class_name: str = Form(None),
    instructor: str = Form(None),
    project_name: str = Form(None),
    user_prompt: str = Form(None),
    advanced_formatting: bool = Form(False),
    design_requirements: str = Form(None),
    knowledge_and_tech: str = Form(None),
    completion: str = Form(None),
    self_statement: str = Form(None),
    textbook: str = Form(None),
    lab: str = Form(None),
    finish_date: str = Form(None),
    generation_mode: str = Form("fusion"),
    file_order: str = Form(None),
    target_pages: str = Form(None),
    multi_round_completion: str = Form('false'),
    cover_template_path: str = Form(None),
    body_template_path: str = Form(None),
    template_id: str = Form(None),
    cover_template: UploadFile = File(None),
    body_template: UploadFile = File(None),
    data_files: list[UploadFile] = File([]),
    db: Session = Depends(get_db)
):
    # 获取当前用户
    user = get_current_user(request, db)
    if not user:
        raise HTTPException(status_code=401, detail="请先登录")
    user_dir = Path(f"uploads/{user.id}")
    user_dir.mkdir(parents=True, exist_ok=True)
    
    # 初始化图片列表
    uploaded_images = []
    
    # 类型转换
    multi_round_completion = (multi_round_completion == 'true')
    try:
        target_pages_int = int(target_pages) if target_pages else None
    except Exception:
        target_pages_int = None
    # 保存上传的资料文件到用户目录
    for file in data_files:
        if file and hasattr(file, 'filename') and file.filename:
            file_path = user_dir / file.filename
            logger.info(f"[上传] 保存文件: {file.filename} -> {file_path}")
            with open(file_path, "wb") as f:
                f.write(await file.read())
            
            # 处理图片文件，生成唯一ID和描述
            if file.filename.lower().endswith(('.jpg', '.jpeg', '.png', '.gif', '.bmp')):
                # 立即复制到 report_images/用户ID/
                import shutil
                dst_path = Path(f"report_images/{user.id}") / file.filename
                shutil.copy(file_path, dst_path)
                try:
                    img_id = f"img_{len(uploaded_images) + 1}"
                    ext = file.filename.split('.')[-1].lower()
                    logger.info(f"[图片处理] 开始处理图片: {file.filename}, ID: {img_id}, 扩展名: {ext}")
                    description = await get_image_description(str(file_path), ext)
                    uploaded_images.append({
                        'id': img_id,
                        'filename': file.filename,
                        'description': description,
                        'filepath': f"{user.id}/{file.filename}",
                        'webpath': f"report_images/{user.id}/{file.filename}"
                    })
                    logger.info(f"[图片处理] 完成: {file.filename}, ID: {img_id}, 描述: {description}")
                except Exception as e:
                    logger.warning(f"[图片处理] 失败: {file.filename}, 错误: {e}")
                    img_id = f"img_{len(uploaded_images) + 1}"
                    uploaded_images.append({
                        'id': img_id,
                        'filename': file.filename,
                        'description': f"图片: {file.filename}",
                        'filepath': f"{user.id}/{file.filename}",
                        'webpath': f"report_images/{user.id}/{file.filename}"
                    })
    # 只遍历该用户目录下的文件
    documents = []
    cover_template_file = None
    body_template_file = None
    data_files_list = []
    for file_path in user_dir.glob("*"):
        logger.info(f"[目录遍历] 检查文件: {file_path}")
        if file_path.name.startswith("cover_template_"):
            cover_template_file = str(file_path)
            logger.info(f"[模板识别] 识别为封面模板: {file_path}")
        elif file_path.name.startswith("body_template_"):
            body_template_file = str(file_path)
            logger.info(f"[模板识别] 识别为正文模板: {file_path}")
        else:
            data_files_list.append(file_path)
            logger.info(f"[资料识别] 识别为资料文件: {file_path}")
    # 遍历uploads目录后，若有模板路径参数则直接使用
    if body_template_path:
        body_template_file = body_template_path
        logger.info(f"[模板参数] 使用传入的正文模板路径: {body_template_path}")
    if cover_template_path:
        cover_template_file = cover_template_path
        logger.info(f"[模板参数] 使用传入的封面模板路径: {cover_template_path}")
    # 1. 自动将uploads目录下所有文档入库（只处理文本文件）
    for file_path in data_files_list:
        logger.info(f"[入库] 处理文档: {file_path}")
        if file_path.is_file():
            fname = file_path.name.lower()
            if any(k in fname for k in ["封面", "cover", "title"]):
                cover_template_file = str(file_path)
                logger.info(f"[入库] 识别为封面模板: {file_path}")
            elif any(k in fname for k in ["正文", "body", "content", "template"]):
                body_template_file = str(file_path)
                logger.info(f"[入库] 识别为正文模板: {file_path}")
            else:
                if file_path.suffix.lower() in ['.docx', '.doc']:
                    word_content = read_word_document_content(file_path)
                    if word_content:
                        documents.append({
                            "content": word_content,
                            "source": str(file_path),
                            "type": file_path.suffix
                        })
                        logger.info(f"[入库] Word文档已入库: {file_path}")
                elif is_text_file(file_path):
                    content = read_text_file_content(file_path)
                    if content:
                        documents.append({
                            "content": content,
                            "source": str(file_path),
                            "type": file_path.suffix
                        })
                        logger.info(f"[入库] 文本文件已入库: {file_path}")
    # ... existing code ...
    logger.info(f"[报告生成] 模式: {generation_mode}, 文档数: {len(documents)}，图片数: {len(uploaded_images)}")
    # ... existing code ...
    if 'report_body' in locals():
        logger.info(f"[报告生成] 报告正文长度: {len(report_body)} 字")
    else:
        logger.warning("[报告生成] report_body 未生成，无法输出长度")
    # ... existing code ...
    logger.info(f"[报告生成] 返回images字段: {uploaded_images}")
    
    # 新增：区分模式下按file_order排序
    if generation_mode == "separate" and file_order:
        try:
            order_list = json.loads(file_order)
            logger.info(f"收到文件顺序: {order_list}")
            logger.info(f"排序前文档数量: {len(documents)}")
            
            # 创建文件名到文档的映射
            doc_map = {}
            for doc in documents:
                file_name = Path(doc["source"]).name
                doc_map[file_name] = doc
                logger.info(f"文档映射: {file_name} -> {doc['source']}")
            
            # 按顺序重新排列文档
            ordered_documents = []
            for name in order_list:
                if name in doc_map:
                    ordered_documents.append(doc_map[name])
                    logger.info(f"按顺序添加文档: {name}")
                else:
                    logger.warning(f"文件顺序中的文件未找到: {name}")
            
            # 添加未在顺序中的文档
            for doc in documents:
                file_name = Path(doc["source"]).name
                if file_name not in order_list:
                    ordered_documents.append(doc)
                    logger.info(f"添加未排序文档: {file_name}")
            
            documents = ordered_documents
            logger.info(f"排序后文档数量: {len(documents)}")
        except Exception as e:
            logger.warning(f"文件顺序解析失败: {e}")
            logger.warning(traceback.format_exc())
    
    if documents:
        chunk_count = 0
        for doc in documents:
            for chunk in split_long_text(doc["content"]):
                if chunk.strip():
                    await ai_service.add_documents_to_vectorstore([{
                        "content": chunk,
                        "source": doc["source"],
                        "type": doc["type"]
                    }])
                    chunk_count += 1
        logger.info(f"已自动入库 {chunk_count} 个分段文档")
    else:
        logger.warning("uploads目录中没有找到可读的文档")
    
    # 2. 根据生成模式选择不同的报告生成策略
    if generation_mode == "separate":
        report_body = await _generate_separate_report(query, documents, user_prompt, advanced_formatting, uploaded_images)
        logger.info(f"区分模式初步整合后内容长度: {len(report_body)} 字")
        if multi_round_completion and target_pages_int:
            logger.info("准备进入多轮补全分支（区分模式）...（对整合后整体内容补全）")
            report_body = await generate_report_to_target_pages(query, report_body, target_pages_int)
        context_for_short = report_body
    else:
        # 融合模式：原有的生成逻辑
        similar_docs = await ai_service.search_similar_documents(query, top_k=5)
        context = [doc["content"][:2000] for doc in similar_docs]
        context_text = "\n\n".join(context)
        MAX_CONTEXT_LEN = 15000
        if len(context_text) > MAX_CONTEXT_LEN:
            context_text = context_text[:MAX_CONTEXT_LEN]
        context = [context_text]
        if not context or not context[0].strip():
            logger.warning("未找到相关文档，使用空上下文")
            context = ["暂无相关文档信息"]
        # 使用多轮补全+自动扩写（受控于multi_round_completion）
        if multi_round_completion and target_pages_int:
            logger.info("准备进入多轮补全分支（融合模式）...")
            report_body = await generate_report_to_target_pages(query, context_text, target_pages_int)
        else:
            final_prompt = build_prompt(query, context_text, user_prompt, advanced_formatting, target_pages_int)
            logger.info(f"使用动态Prompt生成报告，长度: {len(final_prompt)}")
            report_body = await ai_service.generate_report_with_prompt(query, final_prompt, target_pages_int)
        context_for_short = context_text

    # 新增：为设计要求、所用知识与技术、完成情况、自我说明生成简洁内容
    def short_prompt(title):
        return f"""请根据本次实训资料内容，简要生成2-3行的{title}，要求简洁明了。

                        【资料内容】：
                        {context_for_short}
                    """
    design_requirements = design_requirements or (await ai_service.generate_report_with_prompt(query, short_prompt("设计要求"), target_pages_int))
    knowledge_and_tech = knowledge_and_tech or (await ai_service.generate_report_with_prompt(query, short_prompt("所用知识与技术"), target_pages_int))
    completion = completion or (await ai_service.generate_report_with_prompt(query, short_prompt("完成情况"), target_pages_int))
    self_statement = self_statement or (await ai_service.generate_report_with_prompt(query, short_prompt("自我说明"), target_pages_int))

    # 4. 渲染封面和正文模板
    from docxtpl import DocxTemplate
    from docx import Document
    from docxcompose.composer import Composer
    from datetime import datetime
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    # --- 自动修复 name 被覆盖问题 ---
    # 1. 在处理资料文件的循环前，备份表单字段name，防止被覆盖
    name_from_form = name
    context_dict = {
        "name": name_from_form or "",
        "student_id": student_id or "",
        "class_name": class_name or "",
        "instructor": instructor or "",
        "project_name": project_name or "",
        "textbook": textbook or "",
        "lab": lab or "",
        "finish_date": finish_date or "",
        "design_requirements": design_requirements or "",
        "knowledge_and_tech": knowledge_and_tech or "",
        "completion": completion or "",
        "self_statement": self_statement or "",
        "report_body": "{{report_body}}"  # 正文模板先保留占位
    }
    cover_docx = None
    body_docx = None
    if cover_template_file:
        tpl = DocxTemplate(cover_template_file)
        tpl.render(context_dict)
        cover_docx = f"temp/cover_{timestamp}.docx"
        tpl.save(cover_docx)
    if body_template_file:
        tpl = DocxTemplate(body_template_file)
        tpl.render(context_dict)
        body_docx = f"temp/body_{timestamp}.docx"
        tpl.save(body_docx)
    # 5. 用python-docx将AI正文插入正文模板
    if body_docx:
        doc = Document(body_docx)
        target_cell = find_placeholder_cell(doc)
        if target_cell:
            insert_structured_content_to_cell(doc, target_cell, report_body, uploaded_images=uploaded_images)
        else:
            para = doc.add_paragraph()
            insert_structured_content_to_cell(doc, para, report_body, uploaded_images=uploaded_images)
        doc.save(body_docx)
    # 6. 合并封面和正文
    final_docx = None
    if cover_docx and body_docx:
        cover_doc = Document(cover_docx)
        body_doc = Document(body_docx)
        composer = Composer(cover_doc)
        composer.append(body_doc)
        final_docx = f"temp/report_{timestamp}.docx"
        composer.save(final_docx)
        report_filename = f"report_{timestamp}.docx"
        report_path = final_docx
    elif body_docx:
        report_filename = f"report_{timestamp}.docx"
        report_path = body_docx
    else:
        raise Exception("未上传正文模板，无法生成报告")
    # 7. 清理临时文件
    if cover_docx and os.path.exists(cover_docx):
        os.remove(cover_docx)
    if body_docx and os.path.exists(body_docx) and (not final_docx or body_docx != final_docx):
        os.remove(body_docx)
    logger.info(f"报告生成成功: {report_filename}")
    # 8. 生成后清理uploads和向量库，防止历史内容混合
    for file in user_dir.glob("*"):
        if file.is_file():
            file.unlink()
    await ai_service.clear_vectorstore()
    user = get_current_user(request, db)
    if not user:
        raise HTTPException(status_code=401, detail="请先登录")
    if user.usage_count is None or user.usage_count < 1:
        return JSONResponse({"success": False, "detail": "使用次数已用完，请充值后再试"})
    user.usage_count -= 1
    db.commit()
    from datetime import datetime
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    report_id = timestamp  # 以时间戳作为报告唯一ID
    report_image_dir = Path(f"report_images/{user.id}/{report_id}")
    report_image_dir.mkdir(parents=True, exist_ok=True)
    # 复制用到的图片到 report_images 目录
    used_images = []
    for img in uploaded_images:
        src_path = Path("uploads") / img['filepath']
        dst_path = report_image_dir / Path(img['filename'])
        if src_path.exists():
            shutil.copy(src_path, dst_path)
            # 更新图片路径为 report_images 路径
            img['webpath'] = f"report_images/{user.id}/{report_id}/{img['filename']}"
            used_images.append(img)
        else:
            img['webpath'] = ''
            used_images.append(img)
    # 替换 report_body 中的图片占位符为 report_images 路径
    def replace_image_placeholders(text, images):
        for img in images:
            if img['webpath']:
                text = text.replace(f"{{{{image:{img['id']}}}}}", f'<img src="/{img["webpath"]}" alt="{img["description"]}" style="max-width:90%;margin:12px auto;display:block;" />')
            else:
                text = text.replace(f"{{{{image:{img['id']}}}}}", f'<span style="color:red">[图片{img["id"]}未找到]</span>')
        return text
    # ... existing code ...
    # 生成报告正文后，替换图片占位符
    report_body = replace_image_placeholders(report_body, used_images)
    # ... existing code ...
    # 生成报告正文后，自动格式修复
    format_fix_prompt = f"""你是一位文档格式检查与修复助手。请对以下报告内容进行格式检查和修正，要求：\n\n1. 标题层级规范\n2. 图片占位符 {{image:img_x}} 必须单独成段，并在下方补充一句图片描述（如有描述信息）\n3. 列表、编号、代码块等符号符合Markdown规范\n4. 删除多余空行和非法符号\n5. 不要改动正文内容，只做格式修正\n\n【报告内容】：\n{report_body}\n"""
    report_body = await ai_service.generate_report_with_prompt(query, format_fix_prompt)
    return {
        "message": "报告生成成功",
        "report": report_body,
        "filename": report_filename,
        "download_url": f"/download/{report_filename}",
        "images": used_images,
        "context_count": len(context) if 'context' in locals() else 0,
        "advanced_formatting": advanced_formatting,
        "prompt_length": len(final_prompt) if 'final_prompt' in locals() else 0,
        "generation_mode": generation_mode
    }

async def _generate_sections_for_file(query: str, file_name: str, content: str, sections_count: int, uploaded_images=None) -> List[str]:
    """为单个文件生成多个段落"""
    try:
        # 将文件内容分段
        chunks = split_long_text(content)
        
        # 构建图片信息
        image_info = ""
        if uploaded_images:
            image_info = "\n\n【可用图片】：\n"
            for img in uploaded_images:
                image_info += f"- {img['id']}: {img['filename']} - {img['description']}\n"
            image_info += "\n请在合适的位置使用图片占位符格式：{{image:img_x}}，其中x为图片编号。\n"
        
        # 为每个分段生成内容段落
        sections = []
        for i in range(min(sections_count, len(chunks))):
            chunk = chunks[i]
            section_prompt = f"""请根据以下资料内容，为实训报告生成第{i+1}个段落。

【文件来源】：{file_name}
【查询主题】：{query}
【相关资料】：
{chunk}{image_info}

【要求】：
1. 基于提供的资料内容，生成多个个自然段落
2. 内容要具体、详实，避免空泛的描述
3. 语言要正式、专业，符合学术报告风格
4. 如果资料中包含技术细节、数据或示例，请适当引用
5. 段落要有明确的主题和逻辑结构
6. 要尽量详细的写出资料中的所有内容
7. **如果有可用图片，请在合适的位置插入图片占位符**

【格式规则】（必须遵守）：
1. 只允许使用以下 Markdown 符号：`#`（一级标题）、`##`（二级标题）、`###`（三级标题）、```（代码块）、`-`（用于非编号列表）。
2. 禁止使用以下内容：
   - 禁止使用 `---` 分割段落（在每个段落的最后）；
   - 禁止使用 `#####`及以后的标题；
   - 禁止使用 HTML 标签（如 `<p>`、`<br>`）；
   - 禁止写"希望本报告对你有所帮助"等客套话。
   - 禁止在报告标题使用如"实训报告：xxx"的格式，直接写报告名
   - 禁止出现：'• **文本内容**' 的格式, '•' 和 **xxx** 只能单独出现，如：• **向量数据库管理** 这种格式
3. 段落编号请使用 `1.`、`2.`、`3.` 的方式，**不要使用 `-` 和 '•' 作为编号列表**。
4. **图片插入格式为：`{{image:img_x}}`，如：`{{image:img_1}}`，必须出现在内容合适位置**。
5. 所有代码必须使用成对的 ``` 包裹，并保持原始缩进。
6. 除特殊说明外，请保持语气正式、中性、信息导向。
7. 在参考资料的基础上拓展内容，丰富报告的内容
8. 在必要的地方尽可能使用允许的符号"""
            
            section_content = await ai_service.generate_report_with_prompt(query, section_prompt)
            sections.append(section_content)
        
        return sections
    except Exception as e:
        logger.error(f"为文件 {file_name} 生成段落失败: {e}")
        return [f"基于 {file_name} 的内容分析：{str(e)}"]

async def _generate_separate_report(query: str, documents: List[Dict], user_prompt: str = None, advanced_formatting: bool = False, uploaded_images=None) -> str:
    """区分模式：为每个资料生成多个段落，保持资料间的明显分隔"""
    try:
        logger.info(f"开始区分模式报告生成，共 {len(documents)} 个文档")
        
        # 为每个文档生成多个段落
        file_sections = {}
        for doc in documents:
            file_name = Path(doc["source"]).name
            logger.info(f"处理文件: {file_name}")
            
            # 为每个文件生成多个段落，传递图片信息
            sections = await _generate_sections_for_file(query, file_name, doc["content"], 3, uploaded_images)
            file_sections[file_name] = sections
            logger.info(f"为 {file_name} 生成了 {len(sections)} 个段落")
        
        # 简单拼接所有段落，不进行融合
        combined_report = await _combine_sections_without_fusion(query, file_sections, user_prompt, advanced_formatting)
        
        return combined_report
        
    except Exception as e:
        logger.error(f"区分模式报告生成失败: {e}")
        return f"报告生成过程中出现错误: {str(e)}"

async def _combine_sections_without_fusion(query: str, file_sections: Dict[str, List[str]], user_prompt: str = None, advanced_formatting: bool = False) -> str:
    """简单拼接所有段落，不进行内容融合"""
    try:
        # 构建报告开头
        report_header = f"""# {query}

## 概述



"""
        
        # 拼接所有文件的段落
        sections_text = ""
        for file_name, sections in file_sections.items():
            sections_text += f"\n\n## {file_name} \n"
            for i, section in enumerate(sections, 1):
                sections_text += f"\n### {i}\n{section}\n"
        
        # 构建报告结尾
        report_footer = f"""


"""
        
        # 组合完整报告
        full_report = report_header + sections_text + report_footer
        
        # 如果用户有自定义要求，在开头添加说明
        if user_prompt and advanced_formatting:
            full_report = f"""# {query}

## 用户自定义要求

{user_prompt}

{full_report[full_report.find('## 概述'):]}"""
        
        return full_report
        
    except Exception as e:
        logger.error(f"拼接段落失败: {e}")
        return f"报告生成过程中出现错误: {str(e)}"

@app.get("/download/{filename}")
async def download_report(filename: str):
    """下载报告"""
    try:
        file_path = f"temp/{filename}"
        if os.path.exists(file_path):
            return FileResponse(file_path, filename=filename)
        else:
            raise HTTPException(status_code=404, detail="文件不存在")
    except Exception as e:
        logger.error(f"下载文件失败: {e}")
        raise HTTPException(status_code=500, detail=f"下载文件失败: {str(e)}")

@app.post("/add_documents")
async def add_documents():
    """添加文档到向量数据库"""
    try:
        # 读取uploads目录中的文件（包括Word文档）
        upload_dir = Path("uploads")
        documents = []
        
        for file_path in upload_dir.glob("*"):
            if file_path.is_file():
                # 处理Word文档
                if file_path.suffix.lower() in ['.docx', '.doc']:
                    # 检查是否包含模板占位符，如果有则跳过
                    word_content = read_word_document_content(file_path)
                    if word_content and "{{" in word_content and "}}" in word_content:
                        logger.info(f"跳过模板文件: {file_path.name}")
                        continue
                    else:
                        # 不是模板文件，作为资料文档处理
                        if word_content:
                            documents.append({
                                "content": word_content,
                                "source": str(file_path),
                                "type": file_path.suffix
                            })
                            logger.info(f"成功读取Word文档: {file_path.name}")
                        else:
                            logger.warning(f"无法读取Word文档内容: {file_path.name}")
                        continue
                
                # 处理文本文件
                if is_text_file(file_path):
                    content = read_text_file_content(file_path)
                    if content:
                        documents.append({
                            "content": content,
                            "source": str(file_path),
                            "type": file_path.suffix
                        })
                        logger.info(f"成功读取文本文件: {file_path.name}")
                    else:
                        logger.warning(f"无法读取文件内容: {file_path.name}")
                else:
                    logger.info(f"跳过非文本文件: {file_path.name}")
        
        if documents:
            await ai_service.add_documents_to_vectorstore(documents)
            return {"message": f"成功添加 {len(documents)} 个文档到向量数据库"}
        else:
            return {"message": "uploads目录中没有找到可读的文档"}
            
    except Exception as e:
        logger.error(f"添加文档失败: {e}")
        raise HTTPException(status_code=500, detail=f"添加文档失败: {str(e)}")

@app.delete("/cleanup/{session_id}")
async def cleanup_session(session_id: str):
    """清理临时文件"""
    try:
        session_dir = f"temp/{session_id}"
        if os.path.exists(session_dir):
            shutil.rmtree(session_dir)
        return {"success": True, "message": "清理完成"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"清理失败: {str(e)}")

@app.post("/delete_file")
async def delete_file(request: Request):
    data = await request.json()
    filename = data.get("filename")
    try:
        file_path = Path("uploads") / filename
        if file_path.exists() and file_path.is_file():
            file_path.unlink()
            logger.info(f"已删除文件: {filename}")
            return {"success": True, "message": f"已删除文件: {filename}"}
        else:
            return {"success": False, "detail": "文件不存在"}
    except Exception as e:
        logger.error(f"删除文件失败: {e}")
        return {"success": False, "detail": f"删除文件失败: {str(e)}"}

def add_rich_text_to_field(doc, field_name, content, image_dir="uploads"):
    """为指定字段添加富文本内容，仅替换占位符部分，其余内容保留"""
    placeholder_text = f"{{{{{field_name}}}}}"
    target_paragraph = None
    
    # 1. 在段落中查找
    for paragraph in doc.paragraphs:
        if placeholder_text in paragraph.text:
            target_paragraph = paragraph
            logger.info(f"在段落中找到{field_name}字段占位符")
            break
    
    # 2. 在表格中查找
    if not target_paragraph:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if placeholder_text in paragraph.text:
                            target_paragraph = paragraph
                            logger.info(f"在表格中找到{field_name}字段占位符")
                            break
                    if target_paragraph:
                        break
                if target_paragraph:
                    break
            if target_paragraph:
                break
    
    if not target_paragraph:
        logger.warning(f"未找到{field_name}字段占位符: {placeholder_text}")
        return False
    
    # 获取字段对应的样式
    field_style = get_style(doc, [field_name, 'Normal'])
    # 备份原内容
    original_text = target_paragraph.text
    # 清空原有run
    target_paragraph.clear()
    if field_style:
        target_paragraph.style = field_style
        logger.info(f"为{field_name}字段应用样式: {field_style}")
    # 按占位符分割
    before, sep, after = original_text.partition(placeholder_text)
    if before:
        target_paragraph.add_run(before)
    # 插入富文本内容
    if '**' in content:
        parts = content.split('**')
        for i, part in enumerate(parts):
            if part.strip():
                if i % 2 == 1:  # 粗体
                    run = target_paragraph.add_run(part)
                    run.bold = True
                    logger.info(f"为{field_name}字段应用粗体: {part}")
                else:
                    target_paragraph.add_run(part)
    else:
        target_paragraph.add_run(content)
    if after:
        target_paragraph.add_run(after)
    logger.info(f"成功为{field_name}字段添加富文本内容（保留前后内容）")
    return True

def find_placeholder_cell(doc, placeholder_text="{{report_body}}"):
    """查找包含占位符的单元格"""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if placeholder_text in cell.text:
                    logger.info(f"在表格单元格中找到占位符: {cell.text[:50]}...")
                    return cell
    logger.warning(f"未找到包含占位符的单元格: {placeholder_text}")
    return None

def insert_structured_content_to_cell(doc, cell_or_para, markdown_text, uploaded_images=None):
    """在单元格或段落中插入结构化内容（doc为Document对象）"""
    if hasattr(cell_or_para, 'text'):  # 单元格
        cell_or_para.text = ""  # 清空原内容
        container = cell_or_para
    else:  # 段落
        container = cell_or_para
    
    # 获取可用的样式
    heading1_style = get_style(doc, ['Heading 1', '标题 1'])
    heading2_style = get_style(doc, ['Heading 2', '标题 2'])
    heading3_style = get_style(doc, ['Heading 3', '标题 3'])
    heading4_style = get_style(doc, ['Heading 4', '标题 4'])
    code_style = get_style(doc, ['Code'])
    normal_style = get_style(doc, ['Normal', '正文'])
    bold_style = get_style(doc, ['bold', 'Bold', 'Strong'])
    
    logger.info(f"可用样式 - 一级标题: {heading1_style}, 二级标题: {heading2_style}, 三级标题: {heading3_style}, 四级标题: {heading4_style}, 代码样式: {code_style}, 正文: {normal_style}, 粗体: {bold_style}")
    
    lines = markdown_text.strip().split('\n')
    in_code_block = False
    code_lines = []

    for line in lines:
        if line.strip().startswith("```"):
            in_code_block = not in_code_block
            if not in_code_block:
                # 代码块结束，渲染代码
                code_text = '\n'.join(code_lines)
                p = container.add_paragraph()
                if code_style:
                    p.style = code_style
                    # 添加文本内容到已设置样式的段落
                    p.add_run(code_text)
                    logger.info(f"使用Code样式添加代码块: {len(code_lines)}行")
                else:
                    # 如果没有Code样式，手动设置字体
                    run = p.add_run(code_text)
                    run.font.name = "Courier New"
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(64, 64, 64)  # 灰色
                    logger.info(f"手动设置代码块样式: {len(code_lines)}行")
                code_lines.clear()
            continue
        if in_code_block:
            code_lines.append(line)
            continue
        # 行内图片占位符替换
        img_inline = re.findall(r'{{image:(img_\d+)}}', line)
        if img_inline and uploaded_images:
            p = container.add_paragraph()
            p.style = normal_style or 'Normal'
            last_idx = 0
            for match in re.finditer(r'{{image:(img_\d+)}}', line):
                start, end = match.span()
                if start > last_idx:
                    p.add_run(line[last_idx:start])
                img_id = match.group(1)
                img_path = find_image_by_id(img_id, uploaded_images)
                logger.info(f"尝试插入图片: {img_id} 路径: {img_path} 存在: {os.path.exists(img_path) if img_path else False}")
                if img_path and os.path.exists(img_path):
                    run = p.add_run()
                    run.add_picture(img_path, width=Inches(4))
                    logger.info(f"成功插入图片: {img_id} -> {img_path}")
                else:
                    p.add_run(f"[图片{img_id}未找到]")
                last_idx = end
            if last_idx < len(line):
                p.add_run(line[last_idx:])
            continue
        elif line.startswith('#'):
            p = container.add_paragraph(line[2:])
            if heading1_style:
                p.style = heading1_style
                logger.info(f"添加一级标题: {line[2:]} (样式: {heading1_style})")
            else:
                # 如果没有样式，手动设置字体
                run = p.runs[0]
                run.bold = True
                run.font.size = Pt(16)
                logger.info(f"添加一级标题: {line[2:]} (手动设置)")
        elif line.startswith('##'):
            p = container.add_paragraph(line[3:])
            if heading2_style:
                p.style = heading2_style
                logger.info(f"添加二级标题: {line[3:]} (样式: {heading2_style})")
            else:
                # 如果没有样式，手动设置字体
                run = p.runs[0]
                run.bold = True
                run.font.size = Pt(14)
                logger.info(f"添加二级标题: {line[3:]} (手动设置)")
        elif line.startswith('###'):
            p = container.add_paragraph(line[4:])
            if heading3_style:
                p.style = heading3_style
                logger.info(f"添加三级标题: {line[4:]} (样式: {heading3_style})")
            else:
                # 如果没有样式，手动设置字体
                run = p.runs[0]
                run.bold = True
                run.font.size = Pt(12)
                logger.info(f"添加三级标题: {line[4:]} (手动设置)")
        elif line.startswith('####'):
            p = container.add_paragraph(line[5:])
            if heading4_style:
                p.style = heading4_style
                logger.info(f"添加四级标题: {line[5:]} (样式: {heading4_style})")
            else:
                # 如果没有样式，手动设置字体
                run = p.runs[0]
                run.bold = True
                run.font.size = Pt(11)
                logger.info(f"添加四级标题: {line[5:]} (手动设置)")
        elif line.startswith('-'):
            p = container.add_paragraph('• ' + line[2:])
            if normal_style:
                p.style = normal_style
            logger.info(f"添加列表项: {line[2:]}")
        else:
            # 处理富文本格式（粗体等）
            if '**' in line:
                # 分割文本，处理粗体部分
                parts = line.split('**')
                p = container.add_paragraph()
                if normal_style:
                    p.style = normal_style
                
                for i, part in enumerate(parts):
                    if part.strip():  # 跳过空字符串
                        if i % 2 == 1:  # 奇数索引是粗体文本
                            run = p.add_run(part)
                            # 优先使用字符样式，如果没有则使用字体属性
                            if bold_style and bold_style != 'Normal':
                                try:
                                    # 检查是否为字符样式
                                    if doc.styles[bold_style].type == WD_STYLE_TYPE.CHARACTER:
                                        run.style = bold_style
                                        logger.info(f"使用字符样式应用粗体: {part}")
                                    else:
                                        # 如果不是字符样式，使用字体属性
                                        run.bold = True
                                        logger.info(f"使用字体属性应用粗体: {part}")
                                except Exception as e:
                                    # 如果样式应用失败，使用字体属性
                                    run.bold = True
                                    logger.warning(f"样式应用失败，使用字体属性: {e}")
                            else:
                                # 没有找到粗体样式，使用字体属性
                                run.bold = True
                                logger.info(f"使用字体属性应用粗体: {part}")
                        else:  # 偶数索引是普通文本
                            p.add_run(part)
                logger.info(f"添加富文本段落: {line[:50]}...")
            else:
                p = container.add_paragraph(line)
                if normal_style:
                    p.style = normal_style
                logger.info(f"添加普通段落: {line[:30]}...")

def build_prompt(query, context_text, user_prompt=None, advanced_formatting=False, target_pages=None):
    """动态构造Prompt"""
    base_rules = """你是一位实训报告自动生成助手，任务是将以下资料整合为一篇结构清晰、语言通顺、格式严格的 Markdown 报告。请严格按照以下格式要求生成内容：

【格式规则】（必须遵守）：
1. 只允许使用以下 Markdown 符号：`#`（一级标题）、`##`（二级标题）、`###`（三级标题）、```（代码块）、`-`（用于非编号列表）。
2. 禁止使用以下内容：
   - 禁止使用 `---` 分割段落（在每个段落的最后）；
   - 禁止使用 `#####`及以后的标题；
   - 禁止使用 HTML 标签（如 `<p>`、`<br>`）；
   - 禁止写"希望本报告对你有所帮助"等客套话。
   - 禁止在报告标题使用如"实训报告：xxx"的格式，直接写报告名
   - 禁止出现：'• **文本内容**' 的格式, '•' 和 **xxx** 只能单独出现
3. 段落编号请使用 `1.`、`2.`、`3.` 的方式，**不要使用 `-` 和 '•' 作为编号列表**。
4. **图片插入格式为：`{{image:img_x}}`，如：`{{image:img_1}}`，必须出现在内容合适位置**。
5. 所有代码必须使用成对的 ``` 包裹，并保持原始缩进。
6. 除特殊说明外，请保持语气正式、中性、信息导向。
7. 在参考资料的基础上拓展内容，丰富报告的内容
8. 在必要的地方尽可能使用允许的符号"""

    # 添加页面控制要求（如果指定了目标页数）
    page_control = ""
    if target_pages and target_pages > 0:
        # 根据页数调整提示的详细程度
        if target_pages <= 3:
            page_control = f"""

【页面控制要求】：
请确保生成的报告内容大约对应 {target_pages} 页Word文档。要求：
- 每页约包含 500-800 个中文字符
- 内容精简，突出重点
- 避免冗余描述，保持简洁明了"""
        elif target_pages <= 6:
            page_control = f"""

【页面控制要求】：
请确保生成的报告内容大约对应 {target_pages} 页Word文档。要求：
- 每页约包含 500-800 个中文字符
- 内容详实，包含必要的技术细节
- 每个章节要有充分的论述和分析"""
        elif target_pages <= 10:
            page_control = f"""

【页面控制要求】：
请确保生成的报告内容大约对应 {target_pages} 页Word文档。要求：
- 每页约包含 500-800 个中文字符
- 内容非常详细，包含深入的技术分析
- 每个章节都要有充分的论述、案例分析和详细说明
- 可以包含多个子章节和详细的技术实现过程"""
        else:
            page_control = f"""

【页面控制要求】：
请确保生成的报告内容大约对应 {target_pages} 页Word文档。要求：
- 每页约包含 500-800 个中文字符
- 内容极其详细，包含全面的技术分析
- 每个章节都要有充分的论述、案例分析、详细说明和深入探讨
- 必须包含多个子章节、详细的技术实现过程、问题分析、解决方案等
- 可以增加更多的技术细节、实验数据、对比分析等内容
- 确保内容充实，避免空洞的描述"""

    # 如果未启用高级格式自由，则强制使用规则
    if not advanced_formatting:
        prompt = f"""{base_rules}{page_control}

查询问题：{query}

相关资料（供参考）：
{context_text}

请根据内容撰写包含以下部分的报告：
1. 概述
2. 详细分析（可以多小节）
3. 结论与建议
4. 如有必要，可增加技术过程、图示描述等补充部分
5. 尽量避免增加全新内容，而是在原有内容的细节，技术实现等方面做补充

最终请输出一段**结构良好、格式准确、内容严谨的 Markdown 报告**，无需说明"报告结束"等语句。
"""
    else:
        # 启用高级格式设置，用户提示词追加到默认规则后面
        custom_part = user_prompt if user_prompt else ""
        prompt = f"""{base_rules}{page_control}

{custom_part}

查询问题：{query}

上下文信息：
{context_text}
"""
    return prompt

MAX_EMBEDDING_LEN = 8192

def split_long_text(text, max_len=MAX_EMBEDDING_LEN):
    """将长文本按max_len分段"""
    return [text[i:i+max_len] for i in range(0, len(text), max_len)]

# 数据库模型
class User(Base):
    __tablename__ = "users"
    
    id = Column(Integer, primary_key=True, index=True)
    username = Column(String, unique=True, index=True)
    email = Column(String, unique=True, index=True)
    hashed_password = Column(String)
    name = Column(String)
    student_id = Column(String)
    class_name = Column(String)
    created_at = Column(DateTime(timezone=True), server_default=func.now())
    is_active = Column(Boolean, default=True)
    usage_count = Column(Integer, default=20)   # ← 必须加上这一行！

    def __getitem__(self, item):
        return getattr(self, item)

class Template(Base):
    __tablename__ = "templates"
    
    id = Column(Integer, primary_key=True, index=True)
    user_id = Column(Integer, ForeignKey("users.id"))
    name = Column(String, index=True)
    cover_template_path = Column(String)
    body_template_path = Column(String)
    is_example = Column(Boolean, default=False)
    created_at = Column(DateTime(timezone=True), server_default=func.now())

# 订单表
class Order(Base):
    __tablename__ = "orders"
    id = Column(Integer, primary_key=True, index=True)
    user_id = Column(Integer, ForeignKey("users.id"))
    out_trade_no = Column(String, unique=True, index=True)
    amount = Column(String)
    sku_name = Column(String)
    sku_count = Column(Integer)
    created_at = Column(DateTime, default=datetime.utcnow)
    user = relationship("User")

# 创建数据库表
Base.metadata.create_all(bind=engine)

# 密码加密
try:
    pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")
except Exception as e:
    logger.warning(f"bcrypt初始化失败，使用备用方案: {e}")
    # 备用方案：使用简单的哈希
    import hashlib
    pwd_context = None
    
    def get_password_hash(password: str) -> str:
        return hashlib.sha256(password.encode()).hexdigest()
    
    def verify_password(plain_password: str, hashed_password: str) -> bool:
        return hashlib.sha256(plain_password.encode()).hexdigest() == hashed_password
else:
    # 认证相关函数
    def get_password_hash(password: str) -> str:
        return pwd_context.hash(password)

    def verify_password(plain_password: str, hashed_password: str) -> bool:
        return pwd_context.verify(plain_password, hashed_password)

# 自动创建管理员用户
def init_admin_user(db: Session):
    """初始化管理员用户"""
    admin = db.query(User).filter(User.username == 'admin').first()
    if not admin:
        hashed_password = get_password_hash('admin123')
        admin = User(
            username='admin',
            email='admin@example.com',
            hashed_password=hashed_password,
            name='系统管理员',
            student_id='ADMIN001',
            class_name='管理员',
            usage_count=999999  # 管理员无限制使用
        )
        db.add(admin)
        db.commit()
        db.refresh(admin)
        logger.info("管理员用户创建成功")
    return admin

# 初始化管理员用户
with SessionLocal() as db:
    init_admin_user(db)

def get_current_user(request: Request, db: Session = Depends(get_db)):
    user_id = request.session.get("user_id")
    if user_id is None:
        return None
    user = db.query(User).filter(User.id == user_id).first()
    return user

def require_login(request: Request, db: Session = Depends(get_db)):
    user = get_current_user(request, db)
    if user is None:
        raise HTTPException(status_code=status.HTTP_401_UNAUTHORIZED, detail="请先登录")
    return user

# 认证相关路由
@app.get("/login", response_class=HTMLResponse)
async def login_page(request: Request):
    """登录页面"""
    return templates.TemplateResponse("login.html", {"request": request})

@app.get("/register", response_class=HTMLResponse)
async def register_page(request: Request):
    """注册页面"""
    return templates.TemplateResponse("register.html", {"request": request})

@app.post("/login")
async def login(
    request: Request,
    username: str = Form(...),
    password: str = Form(...),
    db: Session = Depends(get_db)
):
    """用户登录"""
    user = db.query(User).filter(User.username == username).first()
    if not user or not verify_password(password, user.hashed_password):
        raise HTTPException(status_code=400, detail="用户名或密码错误")
    
    request.session["user_id"] = user.id
    return RedirectResponse(url="/", status_code=303)

@app.post("/register")
async def register(
    request: Request,
    username: str = Form(...),
    email: str = Form(...),
    password: str = Form(...),
    name: str = Form(...),
    student_id: str = Form(...),
    class_name: str = Form(...),
    db: Session = Depends(get_db)
):
    """用户注册"""
    # 检查用户名是否已存在
    existing_user = db.query(User).filter(User.username == username).first()
    if existing_user:
        raise HTTPException(status_code=400, detail="用户名已存在")
    
    # 检查邮箱是否已存在
    existing_email = db.query(User).filter(User.email == email).first()
    if existing_email:
        raise HTTPException(status_code=400, detail="邮箱已存在")
    
    # 创建新用户
    hashed_password = get_password_hash(password)
    user = User(
        username=username,
        email=email,
        hashed_password=hashed_password,
        name=name,
        student_id=student_id,
        class_name=class_name,
        usage_count=20  # 默认20次
    )
    db.add(user)
    db.commit()
    db.refresh(user)
    
    request.session["user_id"] = user.id
    return RedirectResponse(url="/", status_code=303)

@app.get("/logout")
async def logout(request: Request):
    """用户登出"""
    request.session.clear()
    return RedirectResponse(url="/login", status_code=303)

@app.get("/user/profile")
async def get_user_profile(request: Request, db: Session = Depends(get_db)):
    """获取用户信息"""
    user = get_current_user(request, db)
    if not user:
        raise HTTPException(status_code=401, detail="请先登录")
    return {
        "id": user.id,
        "username": user.username,
        "email": user.email,
        "name": user.name,
        "student_id": user.student_id,
        "class_name": user.class_name,
        "usage_count": getattr(user, "usage_count", None)  # 新增，兼容无此字段情况
    }
@app.get("/user/admin_id")
async def get_admin_id(db: Session = Depends(get_db)):
    admin = db.query(User).filter(User.username == 'admin').first()
    if not admin:
        raise HTTPException(status_code=404, detail="管理员不存在")
    return {"admin_id": admin.id}

@app.post("/user/update_profile")
async def update_user_profile(
    request: Request,
    name: str = Form(...),
    student_id: str = Form(...),
    class_name: str = Form(...),
    db: Session = Depends(get_db)
):
    """更新用户信息"""
    user = require_login(request, db)
    
    user.name = name
    user.student_id = student_id
    user.class_name = class_name
    db.commit()
    
    return {"message": "用户信息更新成功"}

# 模板管理路由
@app.get("/templates")
async def get_user_templates(request: Request, db: Session = Depends(get_db)):
    """获取用户模板列表"""
    user = require_login(request, db)
    
    # 获取用户模板
    user_templates = db.query(Template).filter(
        Template.user_id == user.id,
        Template.is_example == False
    ).all()
    
    # 获取示例模板
    example_templates = db.query(Template).filter(Template.is_example == True).all()
    
    return {
        "user_templates": [
            {
                "id": t.id,
                "name": t.name,
                "cover_template_path": t.cover_template_path,
                "body_template_path": t.body_template_path,
                "created_at": t.created_at.isoformat()
            } for t in user_templates
        ],
        "example_templates": [
            {
                "id": t.id,
                "name": t.name,
                "cover_template_path": t.cover_template_path,
                "body_template_path": t.body_template_path,
                "created_at": t.created_at.isoformat()
            } for t in example_templates
        ]
    }

@app.post("/templates/create")
async def create_template(
    request: Request,
    name: str = Form(...),
    cover_template: UploadFile = File(...),
    body_template: UploadFile = File(...),
    db: Session = Depends(get_db)
):
    """创建新模板"""
    user = require_login(request, db)
    
    # 保存模板文件
    user_template_dir = f"user_templates/{user.id}"
    os.makedirs(user_template_dir, exist_ok=True)
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    cover_path = f"{user_template_dir}/cover_{timestamp}_{cover_template.filename}"
    body_path = f"{user_template_dir}/body_{timestamp}_{body_template.filename}"
    
    # 保存封面模板
    async with aiofiles.open(cover_path, 'wb') as f:
        content = await cover_template.read()
        await f.write(content)
    
    # 保存正文模板
    async with aiofiles.open(body_path, 'wb') as f:
        content = await body_template.read()
        await f.write(content)
    
    # 创建模板记录
    template = Template(
        user_id=user.id,
        name=name,
        cover_template_path=cover_path,
        body_template_path=body_path
    )
    db.add(template)
    db.commit()
    db.refresh(template)
    
    return {"message": "模板创建成功", "template_id": template.id}

@app.delete("/templates/{template_id}")
async def delete_template(
    template_id: int,
    request: Request,
    db: Session = Depends(get_db)
):
    """删除模板"""
    user = require_login(request, db)
    
    template = db.query(Template).filter(
        Template.id == template_id,
        Template.user_id == user.id
    ).first()
    
    if not template:
        raise HTTPException(status_code=404, detail="模板不存在")
    
    # 删除文件
    if os.path.exists(template.cover_template_path):
        os.remove(template.cover_template_path)
    if os.path.exists(template.body_template_path):
        os.remove(template.body_template_path)
    
    # 删除数据库记录
    db.delete(template)
    db.commit()
    
    return {"message": "模板删除成功"}

@app.get("/templates/{template_id}/files")
async def get_template_files(
    template_id: int,
    request: Request,
    db: Session = Depends(get_db)
):
    """获取模板文件路径"""
    user = get_current_user(request, db)
    
    template = db.query(Template).filter(Template.id == template_id).first()
    if not template:
        raise HTTPException(status_code=404, detail="模板不存在")
    
    # 检查权限（用户只能访问自己的模板或示例模板）
    if template.user_id != user.id and not template.is_example:
        raise HTTPException(status_code=403, detail="无权限访问此模板")
    
    return {
        "cover_template_path": template.cover_template_path,
        "body_template_path": template.body_template_path
    }

@app.post("/templates/update/{template_id}")
async def update_template(
    template_id: int,
    request: Request,
    name: str = Form(None),
    cover_template: UploadFile = File(None),
    body_template: UploadFile = File(None),
    db: Session = Depends(get_db)
):
    """更新模板（可更换封面/正文或重命名）"""
    user = require_login(request, db)
    template = db.query(Template).filter(
        Template.id == template_id,
        Template.user_id == user.id
    ).first()
    if not template:
        raise HTTPException(status_code=404, detail="模板不存在")

    user_template_dir = f"user_templates/{user.id}"
    os.makedirs(user_template_dir, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    # 更新封面模板
    if cover_template and cover_template.filename:
        # 删除旧文件
        if template.cover_template_path and os.path.exists(template.cover_template_path):
            os.remove(template.cover_template_path)
        cover_path = f"{user_template_dir}/cover_{timestamp}_{cover_template.filename}"
        async with aiofiles.open(cover_path, 'wb') as f:
            content = await cover_template.read()
            await f.write(content)
        template.cover_template_path = cover_path

    # 更新正文模板
    if body_template and body_template.filename:
        if template.body_template_path and os.path.exists(template.body_template_path):
            os.remove(template.body_template_path)
        body_path = f"{user_template_dir}/body_{timestamp}_{body_template.filename}"
        async with aiofiles.open(body_path, 'wb') as f:
            content = await body_template.read()
            await f.write(content)
        template.body_template_path = body_path

    # 更新名称
    if name:
        template.name = name

    db.commit()
    db.refresh(template)
    return {"message": "模板更新成功", "template_id": template.id}

# 初始化示例模板
def init_example_templates(db: Session):
    """初始化示例模板"""
    # 检查是否已有示例模板
    existing = db.query(Template).filter(Template.is_example == True).first()
    if existing:
        return
    
    # 创建示例模板记录（假设示例文件已存在）
    example_cover = "templates/example_cover_template.docx"
    example_body = "templates/example_body_template.docx"
    
    if os.path.exists(example_cover) and os.path.exists(example_body):
        template = Template(
            user_id=0,  # 系统模板
            name="示例模板",
            cover_template_path=example_cover,
            body_template_path=example_body,
            is_example=True
        )
        db.add(template)
        db.commit()

# 在应用启动时初始化示例模板
with SessionLocal() as db:
    init_example_templates(db)

# 管理员校验（简单示例：假设用户名为admin的为管理员）
def require_admin(request: Request, db: Session = Depends(get_db)):
    user = get_current_user(request, db)
    if not user or user.username != 'admin':
        raise HTTPException(status_code=403, detail="无权限")
    return user

# 用户管理接口
@app.get("/admin/users")
async def admin_get_users(request: Request, db: Session = Depends(get_db)):
    require_admin(request, db)
    users = db.query(User).all()
    return [{
        "id": u.id,
        "username": u.username,
        "name": u.name,
        "student_id": u.student_id,
        "class_name": u.class_name,
        "email": u.email,
        "usage_count": getattr(u, 'usage_count', 0)
    } for u in users]

@app.post("/admin/update_usage")
async def admin_update_usage(request: Request, db: Session = Depends(get_db)):
    require_admin(request, db)
    data = await request.json()
    user_id = data.get('user_id')
    usage_count = data.get('usage_count')
    user = db.query(User).filter(User.id == user_id).first()
    if not user:
        return JSONResponse({"success": False, "detail": "用户不存在"})
    user.usage_count = usage_count
    db.commit()
    return {"success": True}

# 爱发电Webhook接口
@app.post("/pay/afdian/webhook")
async def afdian_webhook(request: Request, db: Session = Depends(get_db)):
    data = await request.json()
    logger.info(f"[AFDIAN] 收到Webhook请求: {data}")
    order = data.get('data', {}).get('order', {})
    remark = order.get('remark', '').strip()
    sku_detail = order.get('sku_detail', [])
    logger.info(f"[AFDIAN] 订单remark: {remark}, sku_detail: {sku_detail}")
    if not remark:
        logger.warning("[AFDIAN] 回调remark为空，测试请求或未填写用户名")
        # 平台测试时返回200，实际不做任何处理
        return {"ec": 200, "em": "测试通过（未填写用户名，未处理业务）"}
    user = db.query(User).filter(User.username == remark).first()
    if not user:
        logger.warning(f"[AFDIAN] 未找到用户: {remark}")
        return {"ec": 404, "em": "未找到用户"}
    total_times = 0
    order_no = order.get('out_trade_no', '')
    amount = order.get('total_amount', '')
    now = datetime.now()
    plan_title = order.get('plan_title', '')  # 新增，商品标题
    for sku in sku_detail:
        sku_id = sku.get('sku_id', '')
        name = sku.get('name', '')
        count = sku.get('count', 1)
        logger.info(f"[AFDIAN] 处理SKU: id={sku_id}, name={name}, count={count}")
        if sku_id in SKU_TIMES_MAP:
            total_times += SKU_TIMES_MAP[sku_id] * count
            logger.info(f"[AFDIAN] SKU精确匹配，增加{SKU_TIMES_MAP[sku_id] * count}次")
        else:
            matched = False
            # 先用plan_title匹配
            for key, times in NAME_TIMES_MAP.items():
                if key in plan_title:
                    total_times += times * count
                    logger.info(f"[AFDIAN] plan_title关键字匹配：{key}，增加{times * count}次")
                    matched = True
            # 再用sku name匹配
            if not matched:
                for key, times in NAME_TIMES_MAP.items():
                    if key in name:
                        total_times += times * count
                        logger.info(f"[AFDIAN] SKU名称关键字匹配：{key}，增加{times * count}次")
                        matched = True
            if not matched:
                logger.warning(f"[AFDIAN] 未匹配到SKU或名称关键字: {name} / {plan_title}")
        # 保存订单记录
        db.add(Order(
            user_id=user.id,
            out_trade_no=order_no,
            amount=amount,
            sku_name=name,
            sku_count=count,
            created_at=now
        ))
        try:
            db.commit()
            logger.info(f"[AFDIAN] 数据库提交成功，用户{user.username}当前剩余{user.usage_count}次")
        except Exception as e:
            logger.error(f"[AFDIAN] 数据库提交失败: {e}")
            db.rollback()
    if total_times == 0:
        total_times = 5  # 默认加10次
        logger.info("[AFDIAN] 未匹配到任何SKU或名称关键字，默认加10次")
    user.usage_count = (user.usage_count or 0) + total_times
    db.commit()
    logger.info(f"[AFDIAN] 用户{user.username}充值成功，增加{total_times}次，当前剩余{user.usage_count}次")
    return {"ec": 200, "em": "ok"}

# 后台：查询用户订单
@app.get("/admin/user/{user_id}/orders")
async def admin_get_user_orders(user_id: int, request: Request, db: Session = Depends(get_db)):
    require_admin(request, db)
    orders = db.query(Order).filter(Order.user_id == user_id).order_by(Order.created_at.desc()).all()
    return [{
        "id": o.id,
        "out_trade_no": o.out_trade_no,
        "amount": o.amount,
        "sku_name": o.sku_name,
        "sku_count": o.sku_count,
        "created_at": o.created_at.isoformat()
    } for o in orders]

# 爱发电商品SKU与次数映射（可为空，主要靠名称关键字）
SKU_TIMES_MAP = {
    # "b082342c4aba11ebb5cb52540025c377": 5,  # 示例：如需精确匹配sku_id可加
}
# 商品名称关键字与次数映射（推荐使用，自动识别）
NAME_TIMES_MAP = {
    "5次生成功能": 5,
    "15次生成功能": 15,
    "25次生成功能": 25,
    "35次生成功能": 30
    # 可继续添加其它商品名关键字
}

# 新增：图片描述生成函数
async def get_image_description(file_path, ext="png"):
    client = OpenAI(
        api_key="sk-442562cd6b6b4b2896ebdac8ce8d047e",
        base_url="https://dashscope.aliyuncs.com/compatible-mode/v1",
    )
    with open(file_path, "rb") as image_file:
        base64_image = base64.b64encode(image_file.read()).decode("utf-8")
    completion = client.chat.completions.create(
        model="qwen-vl-plus",
        messages=[
            {
                "role": "user",
                "content": [
                    {
                        "type": "image_url",
                        "image_url": {"url": f"data:image/{ext};base64,{base64_image}"},
                    },
                    {"type": "text", "text": "请用一句话描述这张图片的内容。"},
                ],
            }
        ],
    )
    return completion.choices[0].message.content

@app.post("/messages/send")
async def send_message(
    request: Request,
    content: str = Form(None),
    image: UploadFile = File(None),
    to_user_id: int = Form(...),
    db: Session = Depends(get_db)
):
    user = get_current_user(request, db)
    if not user:
        raise HTTPException(status_code=401, detail="请先登录")
    
    # 检查目标用户是否存在
    target_user = db.query(User).filter(User.id == to_user_id).first()
    if not target_user:
        raise HTTPException(status_code=404, detail="目标用户不存在")
    
    image_path = None
    if image and image.filename:
        ext = image.filename.split('.')[-1].lower()
        filename = f"{user.id}_{int(datetime.now().timestamp())}.{ext}"
        save_path = os.path.join("chat_images", filename)
        async with aiofiles.open(save_path, 'wb') as f:
            content_bytes = await image.read()
            await f.write(content_bytes)
        image_path = save_path
    
    # 确保至少有一个内容（文字或图片）
    if not content and not image_path:
        raise HTTPException(status_code=422, detail="消息内容不能为空")
    
    msg = Message(
        from_user_id=user.id,
        to_user_id=to_user_id,
        content=content or "",
        image_path=image_path,
        timestamp=datetime.now(),
        is_read=False
    )
    db.add(msg)
    db.commit()
    db.refresh(msg)
    return {"success": True, "message_id": msg.id}

@app.get("/messages/list")
async def get_messages(request: Request, db: Session = Depends(get_db), peer_id: int = None, limit: int = 50):
    user = get_current_user(request, db)
    if not user:
        raise HTTPException(status_code=401, detail="请先登录")
    
    # 获取与peer_id的所有消息（双向）
    if peer_id:
        # 检查peer_id是否为有效用户
        peer_user = db.query(User).filter(User.id == peer_id).first()
        if not peer_user:
            raise HTTPException(status_code=404, detail="目标用户不存在")
        
        msgs = db.query(Message).filter(
            or_(
                (Message.from_user_id == user.id) & (Message.to_user_id == peer_id),
                (Message.from_user_id == peer_id) & (Message.to_user_id == user.id)
            )
        ).order_by(Message.timestamp.asc()).limit(limit).all()
    else:
        # 获取所有与管理员的消息
        admin = db.query(User).filter(User.username == 'admin').first()
        if not admin:
            return []
        msgs = db.query(Message).filter(
            or_(
                (Message.from_user_id == user.id) & (Message.to_user_id == admin.id),
                (Message.from_user_id == admin.id) & (Message.to_user_id == user.id)
            )
        ).order_by(Message.timestamp.asc()).limit(limit).all()
    
    return [{
        "id": m.id,
        "from_user_id": m.from_user_id,
        "to_user_id": m.to_user_id,
        "content": m.content,
        "image_url": f"/chat_images/{os.path.basename(m.image_path)}" if m.image_path else None,
        "timestamp": m.timestamp.isoformat(),
        "is_read": m.is_read
    } for m in msgs]

@app.post("/messages/read")
async def mark_messages_read(request: Request, db: Session = Depends(get_db), peer_id: int = Form(...)):
    user = get_current_user(request, db)
    if not user:
        raise HTTPException(status_code=401, detail="请先登录")
    admin = db.query(User).filter(User.username == 'admin').first()
    if not admin:
        return {"success": False}
    msgs = db.query(Message).filter(
        (Message.from_user_id == admin.id) & (Message.to_user_id == user.id) & (Message.is_read == False)
    ).all()
    for m in msgs:
        m.is_read = True
    db.commit()
    return {"success": True, "count": len(msgs)}

@app.post("/admin/reply")
async def admin_reply_message(
    request: Request,
    to_user_id: int = Form(...),
    content: str = Form(None),
    image: UploadFile = File(None),
    db: Session = Depends(get_db)
):
    """管理员回复消息"""
    # 检查是否为管理员
    user = get_current_user(request, db)
    if not user or user.username != 'admin':
        raise HTTPException(status_code=403, detail="只有管理员可以回复消息")
    
    # 检查目标用户是否存在
    target_user = db.query(User).filter(User.id == to_user_id).first()
    if not target_user:
        raise HTTPException(status_code=404, detail="目标用户不存在")
    
    image_path = None
    if image and image.filename:
        ext = image.filename.split('.')[-1].lower()
        filename = f"admin_{int(datetime.now().timestamp())}.{ext}"
        save_path = os.path.join("chat_images", filename)
        async with aiofiles.open(save_path, 'wb') as f:
            content_bytes = await image.read()
            await f.write(content_bytes)
        image_path = save_path
    
    # 确保至少有一个内容（文字或图片）
    if not content and not image_path:
        raise HTTPException(status_code=422, detail="消息内容不能为空")
    
    msg = Message(
        from_user_id=user.id,
        to_user_id=to_user_id,
        content=content or "",
        image_path=image_path,
        timestamp=datetime.now(),
        is_read=False
    )
    db.add(msg)
    db.commit()
    db.refresh(msg)
    return {"success": True, "message_id": msg.id}

@app.get("/admin/unread_messages")
async def get_unread_messages(request: Request, db: Session = Depends(get_db)):
    """获取所有未读消息（管理员用）"""
    # 检查是否为管理员
    user = get_current_user(request, db)
    if not user or user.username != 'admin':
        raise HTTPException(status_code=403, detail="只有管理员可以查看未读消息")
    
    # 获取所有发给管理员的未读消息
    unread_msgs = db.query(Message).filter(
        (Message.to_user_id == user.id) & (Message.is_read == False)
    ).order_by(Message.timestamp.desc()).all()
    
    # 按用户分组
    user_messages = {}
    for msg in unread_msgs:
        from_user = db.query(User).filter(User.id == msg.from_user_id).first()
        if from_user:
            if from_user.id not in user_messages:
                user_messages[from_user.id] = {
                    "user": {
                        "id": from_user.id,
                        "username": from_user.username,
                        "name": from_user.name,
                        "student_id": from_user.student_id,
                        "class_name": from_user.class_name
                    },
                    "messages": []
                }
            user_messages[from_user.id]["messages"].append({
                "id": msg.id,
                "content": msg.content,
                "image_url": f"/chat_images/{os.path.basename(msg.image_path)}" if msg.image_path else None,
                "timestamp": msg.timestamp.isoformat(),
                "is_read": msg.is_read
            })
    
    return list(user_messages.values())

@app.get("/admin/chat_history/{user_id}")
async def get_chat_history_with_user(
    user_id: int,
    request: Request,
    db: Session = Depends(get_db),
    limit: int = 50
):
    """获取与特定用户的聊天记录（管理员用）"""
    # 检查是否为管理员
    admin = get_current_user(request, db)
    if not admin or admin.username != 'admin':
        raise HTTPException(status_code=403, detail="只有管理员可以查看聊天记录")
    
    # 检查目标用户是否存在
    target_user = db.query(User).filter(User.id == user_id).first()
    if not target_user:
        raise HTTPException(status_code=404, detail="目标用户不存在")
    
    # 获取双向消息
    msgs = db.query(Message).filter(
        or_(
            (Message.from_user_id == admin.id) & (Message.to_user_id == user_id),
            (Message.from_user_id == user_id) & (Message.to_user_id == admin.id)
        )
    ).order_by(Message.timestamp.asc()).limit(limit).all()
    
    return [{
        "id": m.id,
        "from_user_id": m.from_user_id,
        "to_user_id": m.to_user_id,
        "content": m.content,
        "image_url": f"/chat_images/{os.path.basename(m.image_path)}" if m.image_path else None,
        "timestamp": m.timestamp.isoformat(),
        "is_read": m.is_read
    } for m in msgs]

@app.get("/admin/messages", response_class=HTMLResponse)
async def admin_messages_page(request: Request, db: Session = Depends(get_db)):
    """管理员消息管理页面"""
    # 检查是否为管理员
    user = get_current_user(request, db)
    if not user or user.username != 'admin':
        raise HTTPException(status_code=403, detail="只有管理员可以访问此页面")
    
    return templates.TemplateResponse("admin_messages.html", {"request": request})

# 聊天图片静态服务
app.mount("/chat_images", StaticFiles(directory="chat_images"), name="chat_images")

if __name__ == "__main__":
    uvicorn.run(
        "app:app",
        host="0.0.0.0",
        port=8000,
        reload=True,
        log_level="info"
    )