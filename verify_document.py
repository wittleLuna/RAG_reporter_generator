#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
验证生成的Word文档内容
"""

import os
import tempfile
from pathlib import Path
from docx import Document
from test_richtext_fix import test_richtext_fix

def verify_document_content(doc_path):
    """验证文档内容"""
    print(f"🔍 验证文档内容: {doc_path}")
    
    doc = Document(doc_path)
    
    print(f"📄 文档段落数: {len(doc.paragraphs)}")
    print(f"📊 文档表格数: {len(doc.tables)}")
    
    # 检查段落内容
    print("\n📝 段落内容:")
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            print(f"  {i+1}. {para.text[:50]}...")
            # 检查样式
            if para.style:
                print(f"     样式: {para.style.name}")
            # 检查粗体
            for run in para.runs:
                if run.bold:
                    print(f"     粗体: {run.text}")
    
    # 检查表格内容
    print("\n📊 表格内容:")
    for i, table in enumerate(doc.tables):
        print(f"  表格 {i+1}:")
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                if cell.text.strip():
                    print(f"    单元格[{row_idx},{col_idx}]: {cell.text[:30]}...")

def main():
    """主函数"""
    print("🧪 开始验证文档内容...")
    
    # 运行测试生成文档
    doc_path = test_richtext_fix()
    
    if doc_path and os.path.exists(doc_path):
        verify_document_content(doc_path)
        print(f"\n✅ 验证完成，文档路径: {doc_path}")
    else:
        print("❌ 未找到生成的文档")

if __name__ == "__main__":
    main() 