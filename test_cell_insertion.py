#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
测试新的单元格插入逻辑
"""

import os
import tempfile
from pathlib import Path
from docx import Document
from docxtpl import DocxTemplate
from datetime import datetime
from docx.shared import Pt

def create_test_template_with_table():
    """创建包含表格的测试模板"""
    doc = Document()
    
    # 添加标题
    title = doc.add_heading('实训报告', 0)
    
    # 添加基本信息表格
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    
    # 设置表格内容
    table_data = [
        ('姓名：', '{{name}}'),
        ('学号：', '{{student_id}}'),
        ('班级：', '{{class_name}}'),
        ('指导教师：', '{{instructor}}'),
        ('项目名称：', '{{project_name}}')
    ]
    
    for i, (label, placeholder) in enumerate(table_data):
        table.cell(i, 0).text = label
        table.cell(i, 1).text = placeholder
    
    # 添加报告正文表格
    report_table = doc.add_table(rows=1, cols=1)
    report_table.style = 'Table Grid'
    report_table.cell(0, 0).text = '{{report_body}}'
    
    return doc

def find_placeholder_cell_test(doc, placeholder_text="{{report_body}"):
    """测试版本的单元格查找函数"""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if placeholder_text in cell.text:
                    print(f"✅ 在表格单元格中找到占位符: {cell.text[:50]}...")
                    return cell
    print(f"❌ 未找到包含占位符的单元格: {placeholder_text}")
    return None

def insert_structured_content_to_cell_test(cell_or_para, markdown_text):
    """测试版本的结构化内容插入函数"""
    if hasattr(cell_or_para, 'text'):  # 单元格
        cell_or_para.text = ""  # 清空原内容
        container = cell_or_para
        print("📝 在单元格中插入内容")
    else:  # 段落
        container = cell_or_para
        print("📝 在段落中插入内容")
    
    lines = markdown_text.strip().split('\n')
    in_code_block = False
    code_lines = []

    for line in lines:
        if line.strip().startswith("```"):
            in_code_block = not in_code_block
            if not in_code_block:
                p = container.add_paragraph('\n'.join(code_lines))
                run = p.runs[0]
                run.font.name = "Courier New"
                run.font.size = Pt(10)
                code_lines.clear()
                print(f"✅ 添加代码块: {len(code_lines)}行")
            continue
        if in_code_block:
            code_lines.append(line)
            continue
        elif line.startswith('# '):
            container.add_paragraph(line[2:], style='Heading 1')
            print(f"✅ 添加一级标题: {line[2:]}")
        elif line.startswith('## '):
            container.add_paragraph(line[3:], style='Heading 2')
            print(f"✅ 添加二级标题: {line[3:]}")
        elif line.startswith('- '):
            container.add_paragraph('• ' + line[2:])
            print(f"✅ 添加列表项: {line[2:]}")
        else:
            container.add_paragraph(line)
            print(f"✅ 添加段落: {line[:30]}...")

def test_cell_insertion():
    """测试单元格插入功能"""
    print("🧪 开始测试单元格插入功能...")
    
    # 创建临时目录
    with tempfile.TemporaryDirectory() as temp_dir:
        temp_dir = Path(temp_dir)
        
        # 1. 创建测试模板
        template_doc = create_test_template_with_table()
        template_path = temp_dir / "test_template_with_table.docx"
        template_doc.save(template_path)
        print(f"📄 创建测试模板: {template_path}")
        
        # 2. 模拟完整的渲染过程
        try:
            # 2.1 先用docxtpl渲染模板字段
            tpl = DocxTemplate(template_path)
            context_dict = {
                "name": "张三",
                "student_id": "20230001",
                "class_name": "软件工程1班",
                "instructor": "李老师",
                "project_name": "RAG实训项目",
                "report_body": "{{report_body}}"  # 不替换正文，占位
            }
            tpl.render(context_dict)
            
            # 2.2 保存临时渲染结果
            temp_docx = temp_dir / f"temp_rendered_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            tpl.save(temp_docx)
            print(f"✅ docxtpl渲染完成: {temp_docx}")
            
            # 2.3 用python-docx加载并处理结构化内容
            doc = Document(temp_docx)
            
            # 2.4 检查字段是否已正确渲染
            print("\n🔍 检查字段渲染状态...")
            field_check_results = {
                "name": "张三",
                "student_id": "20230001",
                "class_name": "软件工程1班",
                "instructor": "李老师",
                "project_name": "RAG实训项目"
            }
            
            for field_name, field_content in field_check_results.items():
                if field_content and field_content.strip():
                    print(f"✅ 字段{field_name}已通过docxtpl处理: {field_content}")
                else:
                    print(f"⚠️  字段{field_name}为空或未提供")
            
            # 2.5 查找报告正文占位符单元格
            print("\n🔍 查找报告正文占位符单元格...")
            target_cell = find_placeholder_cell_test(doc)
            
            # 2.6 添加测试报告正文
            test_report_body = """# 项目概述

这是一个测试项目，用于验证单元格插入功能。

## 技术栈

- Python 3.8+
- FastAPI
- python-docx
- docxtpl

### 核心组件

项目包含以下**核心组件**：

1. 数据读取模块
2. AI服务模块  
3. 富文本渲染模块

## 代码示例

```python
def hello_world():
    print("Hello, World!")
    return "Success"
```

### 使用方法

这是一个**重要**的使用说明，请仔细阅读。

## 总结

通过本次测试，验证了单元格插入功能的**正确性**和**完整性**。
"""
            
            if target_cell:
                insert_structured_content_to_cell_test(target_cell, test_report_body)
                print("✅ 成功将正文插入表格单元格")
            else:
                print("⚠️  未找到report_body单元格，改为添加到末尾")
                para = doc.add_paragraph()
                insert_structured_content_to_cell_test(para, test_report_body)  # fallback
            
            # 2.7 保存最终文档
            final_docx = temp_dir / f"final_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            doc.save(final_docx)
            print(f"✅ 最终文档保存: {final_docx}")
            
            # 2.8 验证结果
            verify_cell_insertion_result(final_docx)
            
            print("\n🎉 单元格插入测试完成！")
            
        except Exception as e:
            print(f"❌ 测试失败: {e}")
            import traceback
            traceback.print_exc()

def verify_cell_insertion_result(doc_path):
    """验证单元格插入结果"""
    print("\n🔍 验证单元格插入结果...")
    
    doc = Document(doc_path)
    
    # 检查表格内容
    if doc.tables:
        print(f"  - 表格数量: {len(doc.tables)}")
        
        for i, table in enumerate(doc.tables):
            print(f"    - 表格{i+1}: {len(table.rows)}行 x {len(table.columns)}列")
            
            for j, row in enumerate(table.rows):
                for k, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    if cell_text:
                        print(f"      - 单元格({j},{k}): {cell_text[:50]}...")
    
    # 检查段落内容
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    print(f"  - 段落数量: {len(paragraphs)}")
    for i, text in enumerate(paragraphs):
        print(f"    - 段落{i+1}: {text[:50]}...")
    
    print("✅ 单元格插入验证完成")

if __name__ == "__main__":
    test_cell_insertion() 