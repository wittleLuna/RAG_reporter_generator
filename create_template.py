#!/usr/bin/env python3
"""
创建示例Word模板的脚本
"""

import os
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Inches

def create_sample_template():
    """创建示例Word模板"""
    try:
        # 创建一个新的Word文档
        doc = Document()
        
        # 添加标题
        title = doc.add_heading('实训报告', 0)
        title.alignment = 1  # 居中对齐
        
        # 添加基本信息表格
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Table Grid'
        
        # 设置表格内容
        table.cell(0, 0).text = '姓名'
        table.cell(0, 1).text = '{{name}}'
        table.cell(1, 0).text = '学号'
        table.cell(1, 1).text = '{{student_id}}'
        table.cell(2, 0).text = '班级'
        table.cell(2, 1).text = '{{class_name}}'
        table.cell(3, 0).text = '指导教师'
        table.cell(3, 1).text = '{{instructor}}'
        table.cell(4, 0).text = '项目名称'
        table.cell(4, 1).text = '{{project_name}}'
        
        # 添加报告内容标题
        doc.add_heading('报告内容', level=1)
        
        # 添加报告内容占位符
        doc.add_paragraph('{{report_body}}')
        
        # 保存模板
        template_path = 'templates/sample_template.docx'
        os.makedirs('templates', exist_ok=True)
        doc.save(template_path)
        
        print(f"✅ 示例模板创建成功: {template_path}")
        print("模板包含以下占位符:")
        print("- {{name}} - 学生姓名")
        print("- {{student_id}} - 学号")
        print("- {{class_name}} - 班级")
        print("- {{instructor}} - 指导教师")
        print("- {{project_name}} - 项目名称")
        print("- {{report_body}} - 报告内容")
        
        return template_path
        
    except ImportError as e:
        print(f"❌ 缺少依赖: {e}")
        print("请安装: pip install python-docx docxtpl")
        return None
    except Exception as e:
        print(f"❌ 创建模板失败: {e}")
        return None

def test_template():
    """测试模板"""
    try:
        template_path = 'templates/sample_template.docx'
        if not os.path.exists(template_path):
            print("模板文件不存在，先创建模板...")
            template_path = create_sample_template()
            if not template_path:
                return False
        
        # 测试模板渲染
        doc = DocxTemplate(template_path)
        
        context = {
            "name": "张三",
            "student_id": "2021001",
            "class_name": "计算机科学1班",
            "instructor": "李老师",
            "project_name": "Web开发实训",
            "report_body": "这是一个测试报告内容。通过本次实训，我学习了Web开发的基本技术，包括HTML、CSS、JavaScript等。"
        }
        
        doc.render(context)
        
        # 保存测试结果
        test_output = 'templates/test_output.docx'
        doc.save(test_output)
        
        print(f"✅ 模板测试成功: {test_output}")
        return True
        
    except Exception as e:
        print(f"❌ 模板测试失败: {e}")
        return False

if __name__ == "__main__":
    print("=== 创建示例Word模板 ===")
    
    # 创建模板
    template_path = create_sample_template()
    
    if template_path:
        # 测试模板
        print("\n=== 测试模板 ===")
        test_template()
        
        print("\n✅ 模板创建和测试完成！")
        print(f"模板文件: {template_path}")
        print("您可以使用这个模板文件进行测试。")
    else:
        print("❌ 模板创建失败") 