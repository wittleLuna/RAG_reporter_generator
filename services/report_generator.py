import os
import re
from typing import Dict, List, Any
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import shutil

class ReportGenerator:
    """报告生成服务"""
    
    def __init__(self):
        self.supported_image_extensions = {'.png', '.jpg', '.jpeg', '.gif', '.bmp'}
    
    async def generate_report(
        self,
        template_path: str,
        report_content: str,
        user_info: Dict[str, str],
        images: List[str],
        session_dir: str
    ) -> str:
        """生成完整的实训报告"""
        try:
            # 创建输出文件路径
            output_path = os.path.join(session_dir, "generated_report.docx")
            
            # 复制模板文件
            shutil.copy2(template_path, output_path)
            
            # 处理报告内容，分离文本和图片标记
            processed_content = self._process_content_with_images(report_content, images)
            
            # 准备模板变量
            template_vars = {
                "name": user_info["name"],
                "student_id": user_info["student_id"],
                "instructor": user_info["instructor"],
                "project_name": user_info["project_name"],
                "class_name": user_info["class_name"],
                "report_body": processed_content["text"]
            }
            
            # 使用模板生成文档
            doc = DocxTemplate(output_path)
            doc.render(template_vars)
            doc.save(output_path)
            
            # 插入图片
            await self._insert_images(output_path, processed_content["image_placements"])
            
            return output_path
            
        except Exception as e:
            print(f"生成报告失败: {e}")
            raise
    
    def _process_content_with_images(self, content: str, images: List[str]) -> Dict[str, Any]:
        """处理内容，分离文本和图片插入点"""
        # 创建图片映射
        image_map = {}
        for img_path in images:
            filename = os.path.basename(img_path)
            base_name = self._extract_base_name(filename)
            if base_name not in image_map:
                image_map[base_name] = []
            image_map[base_name].append(img_path)
        
        # 按数字排序图片
        for base_name in image_map:
            image_map[base_name].sort(key=lambda x: self._extract_number(os.path.basename(x)))
        
        # 处理内容中的图片标记
        text_parts = []
        image_placements = []
        
        # 按章节分割内容
        sections = re.split(r'(## .+?)(?=## |$)', content, flags=re.DOTALL)
        
        for section in sections:
            if section.strip():
                if section.startswith('## '):
                    # 这是章节标题
                    text_parts.append(section.strip())
                else:
                    # 这是章节内容
                    processed_section = self._process_section_content(section, image_map, image_placements)
                    text_parts.append(processed_section)
        
        return {
            "text": '\n\n'.join(text_parts),
            "image_placements": image_placements
        }
    
    def _process_section_content(self, content: str, image_map: Dict[str, List[str]], image_placements: List[Dict]) -> str:
        """处理章节内容，识别图片插入点"""
        # 查找章节名称
        section_name = self._extract_section_name(content)
        
        # 获取对应的图片列表
        section_images = image_map.get(section_name, [])
        
        # 按段落分割内容
        paragraphs = content.split('\n\n')
        processed_paragraphs = []
        
        for i, paragraph in enumerate(paragraphs):
            processed_paragraphs.append(paragraph)
            
            # 在适当位置插入图片
            if i < len(section_images) and paragraph.strip():
                img_path = section_images[i]
                img_filename = os.path.basename(img_path)
                
                # 添加图片占位符到文本
                processed_paragraphs.append(f"\n[图片: {img_filename}]\n")
                
                # 记录图片插入位置
                image_placements.append({
                    "text_position": len(processed_paragraphs) - 1,
                    "image_path": img_path,
                    "filename": img_filename
                })
        
        return '\n\n'.join(processed_paragraphs)
    
    def _extract_section_name(self, content: str) -> str:
        """从内容中提取章节名称"""
        # 简单的关键词匹配
        keywords = ['openwebui', 'rag', 'llm', 'introduction', 'implementation', 'conclusion']
        
        content_lower = content.lower()
        for keyword in keywords:
            if keyword in content_lower:
                return keyword
        
        return "default"
    
    def _extract_base_name(self, filename: str) -> str:
        """从文件名中提取基础名称"""
        # 移除扩展名和数字
        base_name = re.sub(r'\d+\.(png|jpg|jpeg|gif|bmp)$', '', filename)
        base_name = base_name.rstrip('_').rstrip('-').lower()
        return base_name
    
    def _extract_number(self, filename: str) -> int:
        """从文件名中提取数字"""
        match = re.search(r'(\d+)', filename)
        return int(match.group(1)) if match else 0
    
    async def _insert_images(self, docx_path: str, image_placements: List[Dict]):
        """在Word文档中插入图片"""
        try:
            doc = Document(docx_path)
            
            # 按位置倒序插入图片，避免位置偏移
            image_placements.sort(key=lambda x: x["text_position"], reverse=True)
            
            for placement in image_placements:
                img_path = placement["image_path"]
                text_position = placement["text_position"]
                
                # 找到对应的段落
                if text_position < len(doc.paragraphs):
                    paragraph = doc.paragraphs[text_position]
                    
                    # 在段落后插入图片
                    run = paragraph.add_run()
                    run.add_picture(img_path, width=Inches(5))
                    
                    # 设置图片居中
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 保存文档
            doc.save(docx_path)
            
        except Exception as e:
            print(f"插入图片失败: {e}")
            raise
    
    def create_sample_template(self, output_path: str):
        """创建示例模板文件"""
        doc = Document()
        
        # 添加标题
        title = doc.add_heading('实训报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加基本信息表格
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Table Grid'
        
        # 填充表格
        info_data = [
            ('姓名', '{{name}}'),
            ('学号', '{{student_id}}'),
            ('班级', '{{class_name}}'),
            ('指导老师', '{{instructor}}'),
            ('项目名称', '{{project_name}}')
        ]
        
        for i, (label, value) in enumerate(info_data):
            table.cell(i, 0).text = label
            table.cell(i, 1).text = value
        
        # 添加报告正文占位符
        doc.add_heading('报告正文', level=1)
        doc.add_paragraph('{{report_body}}')
        
        # 保存模板
        doc.save(output_path) 