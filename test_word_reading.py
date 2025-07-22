#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
测试Word文档读取功能
"""

import tempfile
from pathlib import Path
from docx import Document

def read_word_document_content(file_path):
    """读取Word文档内容"""
    try:
        doc = Document(file_path)
        content = []
        
        # 读取段落内容
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                content.append(paragraph.text.strip())
        
        # 读取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_content = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_content.append(cell.text.strip())
                if row_content:
                    content.append(" | ".join(row_content))
        
        return "\n".join(content)
    except Exception as e:
        print(f"无法读取Word文档 {file_path}: {e}")
        return None

def create_test_word_document():
    """创建测试Word文档"""
    doc = Document()
    
    # 添加标题
    title = doc.add_heading('实训报告资料', 0)
    
    # 添加段落内容
    doc.add_paragraph('1. 概述')
    doc.add_paragraph('本次实训旨在通过实际操作和分析，提升对数据处理流程的理解与应用能力。在实训过程中，我们将涉及数据的获取、清洗、分析以及可视化等关键环节。')
    
    doc.add_paragraph('2. 实训目标')
    doc.add_paragraph('通过本次实训，学生应该能够：')
    doc.add_paragraph('- 掌握数据处理的基本流程')
    doc.add_paragraph('- 学会使用相关工具进行数据分析')
    doc.add_paragraph('- 理解数据可视化的原理和方法')
    
    # 添加表格
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    
    # 设置表格内容
    table_data = [
        ('阶段', '内容'),
        ('第一阶段', '数据获取与预处理'),
        ('第二阶段', '数据分析与建模')
    ]
    
    for i, (col1, col2) in enumerate(table_data):
        table.cell(i, 0).text = col1
        table.cell(i, 1).text = col2
    
    return doc

def create_template_word_document():
    """创建模板Word文档"""
    doc = Document()
    
    # 添加标题
    title = doc.add_heading('实训报告模板', 0)
    
    # 添加基本信息表格
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    
    # 设置表格内容
    table_data = [
        ('姓名：', '{{name}}'),
        ('学号：', '{{student_id}}'),
        ('班级：', '{{class_name}}'),
        ('指导教师：', '{{instructor}}'),
        ('项目名称：', '{{project_name}}')
    ]
    
    for i, (label, placeholder) in enumerate(table_data):
        table.cell(i, 0).text = label
        table.cell(i, 1).text = placeholder
    
    # 添加报告正文表格
    report_table = doc.add_table(rows=1, cols=1)
    report_table.style = 'Table Grid'
    report_table.cell(0, 0).text = '{{report_body}}'
    
    return doc

def test_word_reading():
    """测试Word文档读取功能"""
    print("🧪 开始测试Word文档读取功能...")
    
    with tempfile.TemporaryDirectory() as temp_dir:
        temp_dir = Path(temp_dir)
        
        # 1. 创建测试Word文档（资料文档）
        test_doc = create_test_word_document()
        test_doc_path = temp_dir / "test_document.docx"
        test_doc.save(test_doc_path)
        print(f"📄 创建测试Word文档: {test_doc_path}")
        
        # 2. 创建模板Word文档
        template_doc = create_template_word_document()
        template_doc_path = temp_dir / "template_document.docx"
        template_doc.save(template_doc_path)
        print(f"📄 创建模板Word文档: {template_doc_path}")
        
        # 3. 测试读取资料文档
        print("\n🔍 测试读取资料文档...")
        content = read_word_document_content(test_doc_path)
        if content:
            print("✅ 成功读取资料文档内容:")
            print("=" * 50)
            print(content)
            print("=" * 50)
        else:
            print("❌ 读取资料文档失败")
        
        # 4. 测试读取模板文档
        print("\n🔍 测试读取模板文档...")
        template_content = read_word_document_content(template_doc_path)
        if template_content:
            print("✅ 成功读取模板文档内容:")
            print("=" * 50)
            print(template_content)
            print("=" * 50)
            
            # 检查是否包含模板占位符
            if "{{" in template_content and "}}" in template_content:
                print("✅ 模板文档包含占位符，识别为模板文件")
            else:
                print("❌ 模板文档不包含占位符")
        else:
            print("❌ 读取模板文档失败")
        
        print("\n🎉 Word文档读取功能测试完成！")

if __name__ == "__main__":
    test_word_reading() 