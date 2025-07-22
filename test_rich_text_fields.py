#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
测试各个字段的富文本渲染功能
"""

import os
import tempfile
from pathlib import Path
from docx import Document
from docxtpl import DocxTemplate
from datetime import datetime
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, RGBColor

def get_style(doc, style_names):
    """获取存在的样式名，支持多个备选样式名"""
    for style_name in style_names:
        try:
            doc.styles[style_name]
            return style_name
        except KeyError:
            continue
    return None

def create_test_template_with_custom_styles():
    """创建包含自定义样式的测试模板"""
    doc = Document()
    
    # 添加标题
    title = doc.add_heading('实训报告', 0)
    
    # 创建自定义段落样式
    custom_styles = ['name', 'student_id', 'class_name', 'instructor', 'project_name', 'Heading 4']
    
    for style_name in custom_styles:
        try:
            if style_name not in doc.styles:
                # 创建新的段落样式
                new_style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
                # 设置样式属性
                new_style.font.name = '微软雅黑'
                new_style.font.size = Pt(12)
                if style_name == 'name':
                    new_style.font.color.rgb = RGBColor(0, 0, 255)  # 蓝色
                elif style_name == 'student_id':
                    new_style.font.color.rgb = RGBColor(128, 0, 128)  # 紫色
                elif style_name == 'class_name':
                    new_style.font.color.rgb = RGBColor(0, 128, 0)  # 绿色
                elif style_name == 'instructor':
                    new_style.font.color.rgb = RGBColor(255, 0, 0)  # 红色
                elif style_name == 'project_name':
                    new_style.font.color.rgb = RGBColor(255, 165, 0)  # 橙色
                elif style_name == 'Heading 4':
                    new_style.font.bold = True
                    new_style.font.size = Pt(14)
                    new_style.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
                print(f"✅ 成功创建样式: {style_name}")
            else:
                print(f"ℹ️  样式已存在: {style_name}")
        except Exception as e:
            print(f"❌ 创建样式失败 {style_name}: {e}")
    
    # 添加基本信息（使用自定义样式）
    name_para = doc.add_paragraph('姓名：{{name}}')
    name_para.style = doc.styles['name']
    
    student_id_para = doc.add_paragraph('学号：{{student_id}}')
    student_id_para.style = doc.styles['student_id']
    
    class_name_para = doc.add_paragraph('班级：{{class_name}}')
    class_name_para.style = doc.styles['class_name']
    
    instructor_para = doc.add_paragraph('指导教师：{{instructor}}')
    instructor_para.style = doc.styles['instructor']
    
    project_name_para = doc.add_paragraph('项目名称：{{project_name}}')
    project_name_para.style = doc.styles['project_name']
    
    # 添加分隔线
    doc.add_paragraph('=' * 50)
    
    # 添加报告正文占位符
    doc.add_paragraph('{{report_body}}')
    
    return doc

def add_rich_text_to_field(doc, field_name, content, image_dir="uploads"):
    """为指定字段添加富文本内容，仅替换占位符部分，其余内容保留"""
    placeholder_text = f"{{{{{field_name}}}}}"
    target_paragraph = None
    for paragraph in doc.paragraphs:
        if placeholder_text in paragraph.text:
            target_paragraph = paragraph
            break
    if not target_paragraph:
        print(f"⚠️  未找到{field_name}字段占位符: {placeholder_text}")
        # 打印所有段落内容用于调试
        print("📋 当前文档中的所有段落:")
        for i, p in enumerate(doc.paragraphs):
            print(f"  {i}: {p.text}")
        return False
    # 获取字段对应的样式
    field_style = get_style(doc, [field_name, 'Normal'])
    # 备份原内容
    original_text = target_paragraph.text
    # 清空原有run
    target_paragraph.clear()
    if field_style:
        target_paragraph.style = field_style
        print(f"✅ 为{field_name}字段应用样式: {field_style}")
    # 按占位符分割
    before, sep, after = original_text.partition(placeholder_text)
    if before:
        target_paragraph.add_run(before)
    # 插入富文本内容
    if '**' in content:
        parts = content.split('**')
        for i, part in enumerate(parts):
            if part.strip():
                if i % 2 == 1:  # 粗体
                    run = target_paragraph.add_run(part)
                    run.bold = True
                    print(f"✅ 为{field_name}字段应用粗体: {part}")
                else:
                    target_paragraph.add_run(part)
    else:
        target_paragraph.add_run(content)
    if after:
        target_paragraph.add_run(after)
    print(f"✅ 成功为{field_name}字段添加富文本内容（保留前后内容）")
    return True

def process_rendered_markdown_content(doc):
    """处理已渲染的markdown内容，将其转换为真正的Word格式"""
    print("\n🔄 处理已渲染的markdown内容...")
    
    # 定义字段映射（根据段落位置）
    field_mappings = {
        1: ("name", "姓名："),
        2: ("student_id", "学号："),
        3: ("class_name", "班级："),
        4: ("instructor", "指导教师："),
        5: ("project_name", "项目名称：")
    }
    
    for para_index, (field_name, prefix) in field_mappings.items():
        if para_index < len(doc.paragraphs):
            paragraph = doc.paragraphs[para_index]
            original_text = paragraph.text
            
            # 检查是否包含markdown格式
            if '**' in original_text:
                print(f"\n📝 处理字段: {field_name}")
                print(f"   原始内容: {original_text}")
                
                # 获取字段对应的样式
                field_style = get_style(doc, [field_name, 'Normal'])
                
                # 清空原有run
                paragraph.clear()
                
                if field_style:
                    paragraph.style = field_style
                    print(f"✅ 为{field_name}字段应用样式: {field_style}")
                
                # 移除前缀，只处理字段内容部分
                if original_text.startswith(prefix):
                    content_part = original_text[len(prefix):]
                    # 添加前缀
                    paragraph.add_run(prefix)
                    
                    # 处理富文本内容
                    if '**' in content_part:
                        parts = content_part.split('**')
                        for i, part in enumerate(parts):
                            if part.strip():
                                if i % 2 == 1:  # 粗体
                                    run = paragraph.add_run(part)
                                    run.bold = True
                                    print(f"✅ 为{field_name}字段应用粗体: {part}")
                                else:
                                    paragraph.add_run(part)
                    else:
                        paragraph.add_run(content_part)
                    
                    print(f"✅ 成功处理{field_name}字段的富文本内容")
                else:
                    # 如果没有前缀，直接处理整个内容
                    if '**' in original_text:
                        parts = original_text.split('**')
                        for i, part in enumerate(parts):
                            if part.strip():
                                if i % 2 == 1:  # 粗体
                                    run = paragraph.add_run(part)
                                    run.bold = True
                                    print(f"✅ 为{field_name}字段应用粗体: {part}")
                                else:
                                    paragraph.add_run(part)
                    else:
                        paragraph.add_run(original_text)
                    
                    print(f"✅ 成功处理{field_name}字段的富文本内容")

def test_rich_text_fields():
    """测试各个字段的富文本渲染功能"""
    print("🧪 开始测试各个字段的富文本渲染功能...")
    
    # 创建临时目录
    with tempfile.TemporaryDirectory() as temp_dir:
        temp_dir = Path(temp_dir)
        
        # 1. 创建测试模板
        template_doc = create_test_template_with_custom_styles()
        template_path = temp_dir / "test_template_with_fields.docx"
        template_doc.save(template_path)
        print(f"📄 创建测试模板: {template_path}")
        
        # 2. 检查模板中的样式
        print("\n📋 模板中的样式列表:")
        for style in template_doc.styles:
            try:
                style_type = "段落" if style.type == WD_STYLE_TYPE.PARAGRAPH else "字符"
                print(f"  - {style.name} ({style_type})")
            except:
                print(f"  - {style.name} (未知类型)")
        
        # 3. 模拟渲染过程
        try:
            # 3.1 先用docxtpl渲染模板字段（使用富文本内容）
            tpl = DocxTemplate(template_path)
            context_dict = {
                "name": "张三（**优秀学生**）",
                "student_id": "**20230001**（学号）",
                "class_name": "软件工程**1班**",
                "instructor": "**李老师**（副教授）",
                "project_name": "RAG实训**项目**",
                "report_body": "{{report_body}}"  # 保留占位符
            }
            tpl.render(context_dict)
            
            # 3.2 保存临时渲染结果
            temp_docx = temp_dir / f"temp_rendered_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            tpl.save(temp_docx)
            print(f"\n✅ docxtpl渲染完成: {temp_docx}")
            
            # 3.3 用python-docx加载并处理结构化内容
            doc = Document(temp_docx)
            
            # 3.4 检查渲染后的内容
            print("\n📋 渲染后的段落内容:")
            for i, p in enumerate(doc.paragraphs):
                print(f"  {i}: {p.text}")
            
            # 3.5 处理已渲染的markdown内容
            process_rendered_markdown_content(doc)
            
            # 3.6 保存最终文档
            final_docx = temp_dir / f"final_report_with_fields_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            doc.save(final_docx)
            print(f"\n✅ 最终文档保存: {final_docx}")
            
            # 3.7 验证结果
            verify_field_rendering_result(final_docx)
            
            print("\n🎉 字段富文本渲染测试完成！")
            
        except Exception as e:
            print(f"❌ 测试失败: {e}")
            import traceback
            traceback.print_exc()

def verify_field_rendering_result(doc_path):
    """验证字段渲染结果"""
    print("\n🔍 验证字段渲染结果...")
    
    doc = Document(doc_path)
    
    # 检查段落数量
    paragraph_count = len(doc.paragraphs)
    print(f"  - 总段落数: {paragraph_count}")
    
    # 检查各个字段的样式
    field_styles = ['name', 'student_id', 'class_name', 'instructor', 'project_name']
    
    for field_style in field_styles:
        styled_paragraphs = [p for p in doc.paragraphs if p.style.name == field_style]
        print(f"  - {field_style}样式段落数: {len(styled_paragraphs)}")
        for p in styled_paragraphs:
            print(f"    - {p.text}")
    
    # 检查粗体文本
    bold_runs = []
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.bold:
                bold_runs.append(run.text)
    
    print(f"  - 粗体文本数量: {len(bold_runs)}")
    for bold_text in bold_runs:
        print(f"    - {bold_text}")
    
    print("✅ 字段渲染结果验证完成")

if __name__ == "__main__":
    test_rich_text_fields() 