#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
测试字段渲染和报告正文位置修复
"""

import os
import tempfile
from pathlib import Path
from docx import Document
from docxtpl import DocxTemplate
from datetime import datetime
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, RGBColor

def get_style(doc, style_names):
    """获取存在的样式名，支持多个备选样式名"""
    for style_name in style_names:
        try:
            doc.styles[style_name]
            return style_name
        except KeyError:
            continue
    return None

def create_test_template_with_table():
    """创建包含表格的测试模板"""
    doc = Document()
    
    # 添加标题
    title = doc.add_heading('实训报告', 0)
    
    # 添加基本信息表格
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    
    # 设置表格内容
    table_data = [
        ('姓名：', '{{name}}'),
        ('学号：', '{{student_id}}'),
        ('班级：', '{{class_name}}'),
        ('指导教师：', '{{instructor}}'),
        ('项目名称：', '{{project_name}}')
    ]
    
    for i, (label, placeholder) in enumerate(table_data):
        table.cell(i, 0).text = label
        table.cell(i, 1).text = placeholder
    
    # 添加分隔线
    doc.add_paragraph('=' * 50)
    
    # 添加报告正文占位符
    doc.add_paragraph('{{report_body}}')
    
    return doc

def find_placeholder_paragraph_test(doc, placeholder_text="{{report_body}}"):
    """测试版本的占位符查找函数"""
    # 1. 首先在段落中查找
    for paragraph in doc.paragraphs:
        if placeholder_text in paragraph.text:
            print(f"✅ 在段落中找到占位符: {paragraph.text[:50]}...")
            return paragraph
    
    # 2. 在表格中查找
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if placeholder_text in paragraph.text:
                        print(f"✅ 在表格单元格中找到占位符: {paragraph.text[:50]}...")
                        return paragraph
    
    print(f"❌ 未找到占位符: {placeholder_text}")
    return None

def add_rich_text_to_field_test(doc, field_name, content):
    """测试版本的字段富文本处理函数"""
    placeholder_text = f"{{{{{field_name}}}}}"
    target_paragraph = None
    
    # 1. 在段落中查找
    for paragraph in doc.paragraphs:
        if placeholder_text in paragraph.text:
            target_paragraph = paragraph
            print(f"✅ 在段落中找到{field_name}字段占位符")
            break
    
    # 2. 在表格中查找
    if not target_paragraph:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if placeholder_text in paragraph.text:
                            target_paragraph = paragraph
                            print(f"✅ 在表格中找到{field_name}字段占位符")
                            break
                    if target_paragraph:
                        break
                if target_paragraph:
                    break
            if target_paragraph:
                break
    
    if not target_paragraph:
        print(f"❌ 未找到{field_name}字段占位符: {placeholder_text}")
        return False
    
    # 备份原内容
    original_text = target_paragraph.text
    print(f"📝 原始文本: {original_text}")
    
    # 清空原有run
    target_paragraph.clear()
    
    # 按占位符分割
    before, sep, after = original_text.partition(placeholder_text)
    if before:
        target_paragraph.add_run(before)
        print(f"📝 添加前缀: {before}")
    
    # 插入内容
    target_paragraph.add_run(content)
    print(f"📝 插入内容: {content}")
    
    if after:
        target_paragraph.add_run(after)
        print(f"📝 添加后缀: {after}")
    
    print(f"✅ 成功为{field_name}字段添加内容")
    return True

def test_field_rendering():
    """测试字段渲染功能"""
    print("🧪 开始测试字段渲染功能...")
    
    # 创建临时目录
    with tempfile.TemporaryDirectory() as temp_dir:
        temp_dir = Path(temp_dir)
        
        # 1. 创建测试模板
        template_doc = create_test_template_with_table()
        template_path = temp_dir / "test_template_with_table.docx"
        template_doc.save(template_path)
        print(f"📄 创建测试模板: {template_path}")
        
        # 2. 模拟渲染过程
        try:
            # 2.1 先用docxtpl渲染模板字段
            tpl = DocxTemplate(template_path)
            context_dict = {
                "name": "张三",
                "student_id": "20230001",
                "class_name": "软件工程1班",
                "instructor": "李老师",
                "project_name": "RAG实训项目",
                "report_body": "{{report_body}}"  # 保留占位符
            }
            tpl.render(context_dict)
            
            # 2.2 保存临时渲染结果
            temp_docx = temp_dir / f"temp_rendered_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            tpl.save(temp_docx)
            print(f"✅ docxtpl渲染完成: {temp_docx}")
            
            # 2.3 用python-docx加载并处理结构化内容
            doc = Document(temp_docx)
            
            # 2.4 处理各个字段
            field_mappings = {
                "name": "张三",
                "student_id": "20230001",
                "class_name": "软件工程1班",
                "instructor": "李老师",
                "project_name": "RAG实训项目"
            }
            
            print("\n🔧 处理字段...")
            for field_name, field_content in field_mappings.items():
                if field_content and field_content.strip():
                    success = add_rich_text_to_field_test(doc, field_name, field_content)
                    if success:
                        print(f"✅ 成功处理{field_name}字段: {field_content}")
                    else:
                        print(f"❌ 处理{field_name}字段失败")
                else:
                    print(f"⚠️  跳过空字段: {field_name}")
            
            # 2.5 查找报告正文占位符
            print("\n🔍 查找报告正文占位符...")
            target_paragraph = find_placeholder_paragraph_test(doc)
            if not target_paragraph:
                print("⚠️  未找到占位符，在文档末尾添加内容")
                doc.add_paragraph("报告正文：")
                target_paragraph = doc.paragraphs[-1]
            else:
                print("✅ 找到占位符位置")
            
            # 2.6 添加测试内容
            test_content = "这是测试的报告正文内容。"
            target_paragraph.text = test_content
            print(f"✅ 添加报告正文: {test_content}")
            
            # 2.7 保存最终文档
            final_docx = temp_dir / f"final_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            doc.save(final_docx)
            print(f"✅ 最终文档保存: {final_docx}")
            
            # 2.8 验证结果
            verify_field_rendering(final_docx)
            
            print("\n🎉 字段渲染测试完成！")
            
        except Exception as e:
            print(f"❌ 测试失败: {e}")
            import traceback
            traceback.print_exc()

def verify_field_rendering(doc_path):
    """验证字段渲染结果"""
    print("\n🔍 验证字段渲染结果...")
    
    doc = Document(doc_path)
    
    # 检查表格内容
    if doc.tables:
        table = doc.tables[0]
        print(f"  - 表格行数: {len(table.rows)}")
        
        for i, row in enumerate(table.rows):
            if len(row.cells) >= 2:
                label = row.cells[0].text.strip()
                value = row.cells[1].text.strip()
                print(f"    - {label} {value}")
                if value == "":
                    print(f"      ⚠️  字段值为空")
                else:
                    print(f"      ✅ 字段值正常")
    
    # 检查段落内容
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    print(f"  - 段落数量: {len(paragraphs)}")
    for i, text in enumerate(paragraphs):
        print(f"    - 段落{i+1}: {text[:50]}...")
    
    print("✅ 字段渲染验证完成")

if __name__ == "__main__":
    test_field_rendering() 