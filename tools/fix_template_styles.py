#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
模板样式修复脚本
"""

import os
import sys
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, RGBColor

def fix_template_styles(input_path, output_path):
    """修复模板样式"""
    print(f"🔧 修复模板样式: {input_path} -> {output_path}")
    
    try:
        # 加载原始模板
        doc = Document(input_path)
        print(f"✅ 原始模板加载成功")
        
        # 创建新文档
        new_doc = Document()
        print(f"✅ 创建新文档")
        
        # 复制所有样式（避免重复）
        print("📋 复制样式...")
        copied_styles = set()
        
        for style in doc.styles:
            try:
                # 跳过重复的样式
                if style.name in copied_styles:
                    print(f"  ⚠️  跳过重复样式: {style.name}")
                    continue
                
                # 创建新样式
                if style.type == WD_STYLE_TYPE.PARAGRAPH:
                    new_style = new_doc.styles.add_style(style.name, WD_STYLE_TYPE.PARAGRAPH)
                elif style.type == WD_STYLE_TYPE.CHARACTER:
                    new_style = new_doc.styles.add_style(style.name, WD_STYLE_TYPE.CHARACTER)
                elif style.type == WD_STYLE_TYPE.TABLE:
                    new_style = new_doc.styles.add_style(style.name, WD_STYLE_TYPE.TABLE)
                else:
                    continue
                
                # 复制样式属性
                if hasattr(style, 'font') and hasattr(new_style, 'font'):
                    if hasattr(style.font, 'name') and style.font.name:
                        new_style.font.name = style.font.name
                    if hasattr(style.font, 'size') and style.font.size:
                        new_style.font.size = style.font.size
                    if hasattr(style.font, 'bold') and style.font.bold is not None:
                        new_style.font.bold = style.font.bold
                    if hasattr(style.font, 'color') and style.font.color.rgb:
                        new_style.font.color.rgb = style.font.color.rgb
                
                copied_styles.add(style.name)
                print(f"  ✅ 复制样式: {style.name} ({style.type})")
                
            except Exception as e:
                print(f"  ❌ 复制样式失败 {style.name}: {e}")
        
        # 复制文档内容
        print("📝 复制文档内容...")
        for paragraph in doc.paragraphs:
            new_para = new_doc.add_paragraph()
            new_para.text = paragraph.text
            if paragraph.style and paragraph.style.name in copied_styles:
                try:
                    new_para.style = paragraph.style.name
                except:
                    pass
        
        # 保存修复后的模板
        new_doc.save(output_path)
        print(f"✅ 修复后的模板已保存: {output_path}")
        
        return True
        
    except Exception as e:
        print(f"❌ 修复模板失败: {e}")
        return False

def verify_fixed_template(template_path):
    """验证修复后的模板"""
    print(f"\n🔍 验证修复后的模板: {template_path}")
    
    try:
        doc = Document(template_path)
        
        # 检查样式
        print("📋 修复后的样式:")
        style_names = []
        for style in doc.styles:
            try:
                if style.type == WD_STYLE_TYPE.PARAGRAPH:
                    style_type = "段落"
                elif style.type == WD_STYLE_TYPE.CHARACTER:
                    style_type = "字符"
                elif style.type == WD_STYLE_TYPE.TABLE:
                    style_type = "表格"
                else:
                    style_type = str(style.type)
                
                style_names.append(style.name)
                print(f"  - {style.name} ({style_type})")
            except Exception as e:
                print(f"  - {style.name} (类型获取失败: {e})")
        
        # 检查重复
        duplicates = [name for name in set(style_names) if style_names.count(name) > 1]
        if duplicates:
            print(f"⚠️  仍有重复样式: {duplicates}")
        else:
            print("✅ 无重复样式")
        
        # 测试关键样式
        key_styles = ['Heading 1', 'Heading 2', 'Heading 3', 'Heading 4', 'Code', 'Normal', 'bold']
        print(f"\n🎯 关键样式测试:")
        for style_name in key_styles:
            try:
                style = doc.styles[style_name]
                p = doc.add_paragraph(f"测试{style_name}")
                p.style = style
                print(f"  ✅ {style_name}: 应用成功")
            except KeyError:
                print(f"  ❌ {style_name}: 样式不存在")
            except Exception as e:
                print(f"  ❌ {style_name}: 应用失败 - {e}")
        
        return True
        
    except Exception as e:
        print(f"❌ 验证失败: {e}")
        return False

def main():
    """主函数"""
    if len(sys.argv) < 2:
        print("用法: python fix_template_styles.py <原始模板路径> [修复后路径]")
        sys.exit(1)
    
    input_path = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 else input_path.replace('.docx', '_fixed.docx')
    
    print("🚀 模板样式修复工具")
    print("=" * 40)
    
    # 修复模板
    if fix_template_styles(input_path, output_path):
        # 验证修复结果
        verify_fixed_template(output_path)
    
    print("\n✅ 修复完成！")
    print(f"💡 建议使用修复后的模板: {output_path}")

if __name__ == '__main__':
    main() 