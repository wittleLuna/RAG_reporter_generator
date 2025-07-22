#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
简单样式检查脚本
"""

import os
import sys
from docx import Document
from docx.enum.style import WD_STYLE_TYPE

def check_template_styles(template_path):
    """检查模板样式"""
    print(f"🔍 检查模板: {template_path}")
    
    if not os.path.exists(template_path):
        print(f"❌ 模板文件不存在: {template_path}")
        return False
    
    try:
        # 直接加载模板
        doc = Document(template_path)
        print(f"✅ 模板加载成功")
        
        # 检查样式
        print(f"\n📋 模板中的样式:")
        styles_found = []
        for style in doc.styles:
            try:
                if style.type == WD_STYLE_TYPE.PARAGRAPH:
                    style_type = "段落"
                elif style.type == WD_STYLE_TYPE.CHARACTER:
                    style_type = "字符"
                elif style.type == WD_STYLE_TYPE.TABLE:
                    style_type = "表格"
                else:
                    style_type = str(style.type)
                
                styles_found.append((style.name, style_type))
                print(f"  - {style.name} ({style_type})")
            except Exception as e:
                print(f"  - {style.name} (类型获取失败: {e})")
        
        # 检查关键样式
        key_styles = ['Heading 1', 'Heading 2', 'Heading 3', 'Heading 4', 'Code', 'Normal', 'bold']
        print(f"\n🎯 关键样式检查:")
        for style_name in key_styles:
            found = False
            for name, _ in styles_found:
                if name == style_name:
                    found = True
                    break
            print(f"  {style_name}: {'✅' if found else '❌'}")
        
        # 检查重复样式
        style_names = [s[0] for s in styles_found]
        duplicates = [name for name in set(style_names) if style_names.count(name) > 1]
        if duplicates:
            print(f"\n⚠️  发现重复样式: {duplicates}")
        
        return True
        
    except Exception as e:
        print(f"❌ 检查模板失败: {e}")
        return False

def test_style_application(template_path):
    """测试样式应用"""
    print(f"\n🧪 测试样式应用:")
    
    try:
        doc = Document(template_path)
        
        # 测试关键样式
        test_styles = ['Heading 1', 'Heading 2', 'Code', 'Normal']
        for style_name in test_styles:
            try:
                style = doc.styles[style_name]
                p = doc.add_paragraph(f"测试{style_name}样式")
                p.style = style
                print(f"  ✅ {style_name}: 应用成功")
            except KeyError:
                print(f"  ❌ {style_name}: 样式不存在")
            except Exception as e:
                print(f"  ❌ {style_name}: 应用失败 - {e}")
        
        return True
        
    except Exception as e:
        print(f"❌ 测试样式应用失败: {e}")
        return False

def main():
    """主函数"""
    if len(sys.argv) < 2:
        print("用法: python simple_style_check.py <模板文件路径>")
        sys.exit(1)
    
    template_path = sys.argv[1]
    
    print("🚀 简单样式检查工具")
    print("=" * 40)
    
    # 检查模板样式
    if check_template_styles(template_path):
        # 测试样式应用
        test_style_application(template_path)
    
    print("\n✅ 检查完成！")

if __name__ == '__main__':
    main() 