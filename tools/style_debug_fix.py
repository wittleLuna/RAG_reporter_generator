#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
样式调试和自动修复脚本
"""

import os
import sys
import tempfile
from pathlib import Path
from docx import Document
from docxtpl import DocxTemplate
from datetime import datetime
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, RGBColor

def get_style(doc, style_names):
    """获取存在的样式名，支持多个备选样式名"""
    for style_name in style_names:
        try:
            doc.styles[style_name]
            return style_name
        except KeyError:
            continue
    return None

def debug_template_styles(template_path):
    """调试模板样式问题"""
    print(f"🔍 开始调试模板: {template_path}")
    
    if not os.path.exists(template_path):
        print(f"❌ 模板文件不存在: {template_path}")
        return False
    
    try:
        # 1. 直接加载模板检查样式
        print("\n📋 直接加载模板的样式:")
        doc = Document(template_path)
        direct_styles = []
        for style in doc.styles:
            try:
                style_type = "段落" if style.type == WD_STYLE_TYPE.PARAGRAPH else "字符"
                direct_styles.append((style.name, style_type))
                print(f"  - {style.name} ({style_type})")
            except Exception as e:
                print(f"  - {style.name} (类型获取失败: {e})")
        
        # 2. 使用docxtpl加载模板检查样式
        print("\n📋 使用docxtpl加载模板的样式:")
        tpl = DocxTemplate(template_path)
        docxtpl_styles = []
        for style in tpl.docx.styles:
            try:
                style_type = "段落" if style.type == WD_STYLE_TYPE.PARAGRAPH else "字符"
                docxtpl_styles.append((style.name, style_type))
                print(f"  - {style.name} ({style_type})")
            except Exception as e:
                print(f"  - {style.name} (类型获取失败: {e})")
        
        # 3. 检查关键样式
        key_styles = ['Heading 1', 'Heading 2', 'Heading 3', 'Heading 4', 'Code', 'Normal', 'bold']
        print(f"\n🎯 关键样式检查:")
        for style_name in key_styles:
            direct_found = get_style(doc, [style_name])
            docxtpl_found = get_style(tpl.docx, [style_name])
            print(f"  {style_name}: 直接加载={'✅' if direct_found else '❌'}, docxtpl加载={'✅' if docxtpl_found else '❌'}")
        
        # 4. 检查样式差异
        direct_style_names = set([s[0] for s in direct_styles])
        docxtpl_style_names = set([s[0] for s in docxtpl_styles])
        
        if direct_style_names != docxtpl_style_names:
            print(f"\n⚠️  样式差异检测:")
            missing_in_docxtpl = direct_style_names - docxtpl_style_names
            if missing_in_docxtpl:
                print(f"  docxtpl中缺失的样式: {missing_in_docxtpl}")
            
            extra_in_docxtpl = docxtpl_style_names - direct_style_names
            if extra_in_docxtpl:
                print(f"  docxtpl中额外的样式: {extra_in_docxtpl}")
        
        return True
        
    except Exception as e:
        print(f"❌ 调试模板失败: {e}")
        return False

def create_fixed_template(original_path, output_path):
    """创建修复后的模板"""
    print(f"\n🔧 创建修复后的模板: {output_path}")
    
    try:
        # 加载原始模板
        doc = Document(original_path)
        
        # 创建新的文档
        new_doc = Document()
        
        # 复制所有样式
        print("📋 复制样式...")
        for style in doc.styles:
            try:
                if style.name not in new_doc.styles:
                    # 创建新样式
                    if style.type == WD_STYLE_TYPE.PARAGRAPH:
                        new_style = new_doc.styles.add_style(style.name, WD_STYLE_TYPE.PARAGRAPH)
                    elif style.type == WD_STYLE_TYPE.CHARACTER:
                        new_style = new_doc.styles.add_style(style.name, WD_STYLE_TYPE.CHARACTER)
                    else:
                        continue
                    
                    # 复制样式属性
                    if hasattr(style, 'font') and hasattr(new_style, 'font'):
                        if hasattr(style.font, 'name') and style.font.name:
                            new_style.font.name = style.font.name
                        if hasattr(style.font, 'size') and style.font.size:
                            new_style.font.size = style.font.size
                        if hasattr(style.font, 'bold') and style.font.bold is not None:
                            new_style.font.bold = style.font.bold
                        if hasattr(style.font, 'color') and style.font.color.rgb:
                            new_style.font.color.rgb = style.font.color.rgb
                    
                    print(f"  ✅ 复制样式: {style.name}")
                else:
                    print(f"  ℹ️  样式已存在: {style.name}")
            except Exception as e:
                print(f"  ❌ 复制样式失败 {style.name}: {e}")
        
        # 复制文档内容
        print("📝 复制文档内容...")
        for paragraph in doc.paragraphs:
            new_para = new_doc.add_paragraph()
            new_para.text = paragraph.text
            if paragraph.style:
                try:
                    new_para.style = paragraph.style.name
                except:
                    pass
        
        # 保存修复后的模板
        new_doc.save(output_path)
        print(f"✅ 修复后的模板已保存: {output_path}")
        
        return True
        
    except Exception as e:
        print(f"❌ 创建修复模板失败: {e}")
        return False

def test_style_rendering(template_path):
    """测试样式渲染"""
    print(f"\n🧪 测试样式渲染: {template_path}")
    
    try:
        # 创建临时目录
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_dir = Path(temp_dir)
            
            # 1. 使用docxtpl渲染
            tpl = DocxTemplate(template_path)
            context = {
                "name": "测试用户",
                "student_id": "20230001",
                "class_name": "测试班级",
                "instructor": "测试教师",
                "project_name": "测试项目",
                "report_body": "{{report_body}}"
            }
            tpl.render(context)
            
            # 保存渲染结果
            rendered_path = temp_dir / "rendered_template.docx"
            tpl.save(rendered_path)
            print(f"✅ docxtpl渲染完成: {rendered_path}")
            
            # 2. 检查渲染后的样式
            doc = Document(rendered_path)
            print("📋 渲染后的样式:")
            for style in doc.styles:
                try:
                    style_type = "段落" if style.type == WD_STYLE_TYPE.PARAGRAPH else "字符"
                    print(f"  - {style.name} ({style_type})")
                except:
                    print(f"  - {style.name} (未知类型)")
            
            # 3. 测试样式应用
            print("\n🎯 测试样式应用:")
            test_styles = ['Heading 1', 'Heading 2', 'Code', 'Normal']
            for style_name in test_styles:
                found_style = get_style(doc, [style_name])
                if found_style:
                    try:
                        p = doc.add_paragraph(f"测试{style_name}样式")
                        p.style = found_style
                        print(f"  ✅ {style_name}: 应用成功")
                    except Exception as e:
                        print(f"  ❌ {style_name}: 应用失败 - {e}")
                else:
                    print(f"  ❌ {style_name}: 未找到")
            
            # 保存测试结果
            test_result_path = temp_dir / "test_result.docx"
            doc.save(test_result_path)
            print(f"✅ 测试结果已保存: {test_result_path}")
            
            return True
            
    except Exception as e:
        print(f"❌ 测试样式渲染失败: {e}")
        return False

def main():
    """主函数"""
    if len(sys.argv) < 2:
        print("用法: python style_debug_fix.py <模板文件路径> [修复输出路径]")
        sys.exit(1)
    
    template_path = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 else None
    
    print("🚀 样式调试和自动修复工具")
    print("=" * 50)
    
    # 1. 调试模板样式
    if not debug_template_styles(template_path):
        print("❌ 调试失败，退出")
        sys.exit(1)
    
    # 2. 测试样式渲染
    test_style_rendering(template_path)
    
    # 3. 创建修复模板（如果指定了输出路径）
    if output_path:
        create_fixed_template(template_path, output_path)
    
    print("\n✅ 调试完成！")
    print("\n💡 建议:")
    print("1. 如果样式在docxtpl中缺失，尝试重新保存模板")
    print("2. 如果样式应用失败，检查样式类型是否正确")
    print("3. 如果问题持续，使用修复后的模板")

if __name__ == '__main__':
    main() 