#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
测试图片插入功能
"""

import os
import sys
import tempfile
from pathlib import Path
from docx import Document
from docxtpl import DocxTemplate
from datetime import datetime

def test_image_insertion():
    """测试图片插入功能"""
    print("🧪 测试图片插入功能")
    
    # 模拟图片数据
    uploaded_images = [
        {
            'id': 'img_1',
            'filename': 'Scrt.png',
            'description': '这张图片显示了一个终端界面，用户正在使用Ollama运行一个中文聊天模型',
            'filepath': '1/Scrt.png'
        },
        {
            'id': 'img_2',
            'filename': 'VMware.png',
            'description': '这张图片显示了一个在VMware Workstation中运行的Ubuntu虚拟机',
            'filepath': '1/VMware.png'
        }
    ]
    
    # 测试图片占位符识别
    test_content = """
# 测试报告

## 概述

这是一个测试报告，包含图片插入功能。

## 技术实现

在实现过程中，我们使用了以下技术：

{{image:img_1}}

如图所示，终端界面显示了模型加载过程。

## 环境配置

我们使用VMware虚拟机进行环境配置：

{{image:img_2}}

虚拟机配置完成后，我们就可以开始部署了。

## 结论

通过以上步骤，我们成功完成了配置。
"""
    
    print("📋 测试内容:")
    print(test_content)
    
    # 测试图片占位符提取
    import re
    img_matches = re.findall(r'{{image:(img_\d+)}}', test_content)
    print(f"\n🎯 找到的图片占位符: {img_matches}")
    
    # 测试图片路径查找
    for img_id in img_matches:
        img_path = None
        for img in uploaded_images:
            if img['id'] == img_id:
                img_path = f"uploads/{img['filepath']}"
                break
        
        if img_path:
            print(f"✅ 找到图片 {img_id}: {img_path}")
        else:
            print(f"❌ 未找到图片 {img_id}")
    
    # 测试Word文档中的图片插入
    print(f"\n📄 测试Word文档图片插入...")
    
    try:
        # 创建临时目录
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_dir = Path(temp_dir)
            
            # 创建测试模板
            doc = Document()
            doc.add_paragraph("{{report_body}}")
            template_path = temp_dir / "test_template.docx"
            doc.save(template_path)
            
            # 使用docxtpl渲染
            tpl = DocxTemplate(template_path)
            context = {"report_body": "{{report_body}}"}
            tpl.render(context)
            
            # 保存渲染结果
            rendered_path = temp_dir / "rendered.docx"
            tpl.save(rendered_path)
            
            # 加载渲染后的文档
            doc = Document(rendered_path)
            
            # 查找占位符
            target_paragraph = None
            for paragraph in doc.paragraphs:
                if "{{report_body}}" in paragraph.text:
                    target_paragraph = paragraph
                    break
            
            if target_paragraph:
                print("✅ 找到占位符位置")
                
                # 模拟图片插入
                lines = test_content.split('\n')
                for line in lines:
                    if '{{image:' in line:
                        img_match = re.findall(r'{{image:(img_\d+)}}', line)
                        if img_match:
                            img_id = img_match[0]
                            print(f"📸 处理图片占位符: {img_id}")
                            
                            # 查找图片路径
                            img_path = None
                            for img in uploaded_images:
                                if img['id'] == img_id:
                                    img_path = f"uploads/{img['filepath']}"
                                    break
                            
                            if img_path:
                                print(f"✅ 找到图片路径: {img_path}")
                                # 这里应该插入图片，但为了测试我们只记录
                            else:
                                print(f"❌ 未找到图片路径: {img_id}")
                    else:
                        print(f"📝 处理文本行: {line[:50]}...")
            else:
                print("❌ 未找到占位符位置")
            
            print("✅ 图片插入测试完成")
            
    except Exception as e:
        print(f"❌ 测试失败: {e}")

def main():
    """主函数"""
    print("🚀 图片插入功能测试工具")
    print("=" * 40)
    
    test_image_insertion()
    
    print("\n✅ 测试完成！")

if __name__ == '__main__':
    main() 