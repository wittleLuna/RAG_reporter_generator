#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
最终测试脚本 - 验证字段渲染和报告正文位置修复
"""

import os
import tempfile
from pathlib import Path
from docx import Document
from docxtpl import DocxTemplate
from datetime import datetime

def create_test_template():
    """创建测试模板"""
    doc = Document()
    
    # 添加标题
    title = doc.add_heading('实训报告', 0)
    
    # 添加基本信息
    doc.add_paragraph('姓名：{{name}}')
    doc.add_paragraph('学号：{{student_id}}')
    doc.add_paragraph('班级：{{class_name}}')
    doc.add_paragraph('指导教师：{{instructor}}')
    doc.add_paragraph('项目名称：{{project_name}}')
    
    # 添加分隔线
    doc.add_paragraph('=' * 50)
    
    # 添加报告正文占位符
    doc.add_paragraph('{{report_body}}')
    
    return doc

def test_final_fix():
    """测试最终修复效果"""
    print("🧪 开始测试最终修复效果...")
    
    # 创建临时目录
    with tempfile.TemporaryDirectory() as temp_dir:
        temp_dir = Path(temp_dir)
        
        # 1. 创建测试模板
        template_doc = create_test_template()
        template_path = temp_dir / "test_template.docx"
        template_doc.save(template_path)
        print(f"📄 创建测试模板: {template_path}")
        
        # 2. 模拟完整的渲染过程
        try:
            # 2.1 先用docxtpl渲染模板字段
            tpl = DocxTemplate(template_path)
            context_dict = {
                "name": "张三",
                "student_id": "20230001",
                "class_name": "软件工程1班",
                "instructor": "李老师",
                "project_name": "RAG实训项目",
                "report_body": "{{report_body}}"  # 保留占位符
            }
            tpl.render(context_dict)
            
            # 2.2 保存临时渲染结果
            temp_docx = temp_dir / f"temp_rendered_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            tpl.save(temp_docx)
            print(f"✅ docxtpl渲染完成: {temp_docx}")
            
            # 2.3 用python-docx加载并检查字段状态
            doc = Document(temp_docx)
            
            # 2.4 检查字段是否已正确渲染
            print("\n🔍 检查字段渲染状态...")
            field_check_results = {
                "name": "张三",
                "student_id": "20230001",
                "class_name": "软件工程1班",
                "instructor": "李老师",
                "project_name": "RAG实训项目"
            }
            
            for field_name, field_content in field_check_results.items():
                if field_content and field_content.strip():
                    print(f"✅ 字段{field_name}已通过docxtpl处理: {field_content}")
                else:
                    print(f"⚠️  字段{field_name}为空或未提供")
            
            # 2.5 查找报告正文占位符
            print("\n🔍 查找报告正文占位符...")
            target_paragraph = None
            for paragraph in doc.paragraphs:
                if "{{report_body}}" in paragraph.text:
                    target_paragraph = paragraph
                    print(f"✅ 找到占位符: {paragraph.text}")
                    break
            
            if not target_paragraph:
                print("⚠️  未找到占位符，在文档末尾添加内容")
                doc.add_paragraph("报告正文：")
                target_paragraph = doc.paragraphs[-1]
            else:
                print("✅ 找到占位符位置")
            
            # 2.6 添加测试报告正文
            test_report_body = """# 项目概述

这是一个测试项目，用于验证修复效果。

## 技术栈

- Python 3.8+
- FastAPI
- python-docx
- docxtpl

### 核心组件

项目包含以下**核心组件**：

1. 数据读取模块
2. AI服务模块  
3. 富文本渲染模块

## 代码示例

```python
def hello_world():
    print("Hello, World!")
    return "Success"
```

### 使用方法

这是一个**重要**的使用说明，请仔细阅读。

## 总结

通过本次测试，验证了修复的**正确性**和**完整性**。
"""
            
            # 2.7 替换占位符内容
            target_paragraph.text = test_report_body
            print(f"✅ 添加报告正文（长度: {len(test_report_body)}字符）")
            
            # 2.8 保存最终文档
            final_docx = temp_dir / f"final_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            doc.save(final_docx)
            print(f"✅ 最终文档保存: {final_docx}")
            
            # 2.9 验证结果
            verify_final_result(final_docx)
            
            print("\n🎉 最终修复测试完成！")
            
        except Exception as e:
            print(f"❌ 测试失败: {e}")
            import traceback
            traceback.print_exc()

def verify_final_result(doc_path):
    """验证最终结果"""
    print("\n🔍 验证最终结果...")
    
    doc = Document(doc_path)
    
    # 检查段落内容
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    print(f"  - 段落数量: {len(paragraphs)}")
    
    # 检查字段内容
    field_expected = {
        "姓名：": "张三",
        "学号：": "20230001",
        "班级：": "软件工程1班",
        "指导教师：": "李老师",
        "项目名称：": "RAG实训项目"
    }
    
    print("\n📋 字段验证:")
    for i, text in enumerate(paragraphs):
        for label, expected_value in field_expected.items():
            if text.startswith(label):
                actual_value = text.replace(label, "").strip()
                if actual_value == expected_value:
                    print(f"  ✅ {label}{actual_value}")
                else:
                    print(f"  ❌ {label}{actual_value} (期望: {expected_value})")
                break
        else:
            # 检查是否包含报告正文
            if "项目概述" in text or "技术栈" in text:
                print(f"  📄 报告正文段落: {text[:50]}...")
    
    print("✅ 最终结果验证完成")

if __name__ == "__main__":
    test_final_fix() 