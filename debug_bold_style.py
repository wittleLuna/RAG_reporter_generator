#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
调试bold样式问题
"""

import os
import tempfile
from pathlib import Path
from docx import Document
from docxtpl import DocxTemplate
from datetime import datetime
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, RGBColor

def get_style(doc, style_names):
    """获取存在的样式名，支持多个备选样式名"""
    for style_name in style_names:
        try:
            doc.styles[style_name]
            return style_name
        except KeyError:
            continue
    return None

def create_test_template_with_bold():
    """创建包含bold字符样式的测试模板"""
    doc = Document()
    
    # 添加标题
    title = doc.add_heading('测试模板', 0)
    
    # 尝试创建bold字符样式
    try:
        # 检查是否已存在bold样式
        if 'bold' not in doc.styles:
            # 创建新的字符样式
            bold_style = doc.styles.add_style('bold', WD_STYLE_TYPE.CHARACTER)
            bold_style.font.bold = True
            bold_style.font.color.rgb = RGBColor(0, 0, 255)  # 蓝色
            print("✅ 成功创建bold字符样式")
        else:
            print("ℹ️  bold样式已存在")
    except Exception as e:
        print(f"❌ 创建bold样式失败: {e}")
    
    # 添加基本信息
    doc.add_paragraph('姓名：{{name}}')
    doc.add_paragraph('学号：{{student_id}}')
    doc.add_paragraph('班级：{{class_name}}')
    doc.add_paragraph('指导教师：{{instructor}}')
    doc.add_paragraph('项目名称：{{project_name}}')
    
    # 添加分隔线
    doc.add_paragraph('=' * 50)
    
    # 添加报告正文占位符
    doc.add_paragraph('{{report_body}}')
    
    return doc

def debug_bold_style():
    """调试bold样式问题"""
    print("🔍 开始调试bold样式问题...")
    
    # 创建临时目录
    with tempfile.TemporaryDirectory() as temp_dir:
        temp_dir = Path(temp_dir)
        
        # 1. 创建测试模板
        template_doc = create_test_template_with_bold()
        template_path = temp_dir / "test_template_with_bold.docx"
        template_doc.save(template_path)
        print(f"📄 创建测试模板: {template_path}")
        
        # 2. 检查模板中的样式
        print("\n📋 模板中的样式列表:")
        for style in template_doc.styles:
            try:
                style_type = "段落" if style.type == WD_STYLE_TYPE.PARAGRAPH else "字符"
                print(f"  - {style.name} ({style_type})")
            except:
                print(f"  - {style.name} (未知类型)")
        
        # 3. 检查bold样式
        bold_style = get_style(template_doc, ['bold', 'Bold', 'Strong'])
        if bold_style:
            try:
                style_obj = template_doc.styles[bold_style]
                style_type = "段落" if style_obj.type == WD_STYLE_TYPE.PARAGRAPH else "字符"
                print(f"\n✅ 找到bold样式: {bold_style} ({style_type})")
                
                # 检查样式属性
                if hasattr(style_obj, 'font'):
                    print(f"  - 字体加粗: {style_obj.font.bold}")
                    print(f"  - 字体颜色: {style_obj.font.color.rgb}")
            except Exception as e:
                print(f"❌ 检查bold样式失败: {e}")
        else:
            print("\n❌ 未找到bold样式")
        
        # 4. 模拟渲染过程
        try:
            # 4.1 先用docxtpl渲染模板字段
            tpl = DocxTemplate(template_path)
            context_dict = {
                "name": "张三",
                "student_id": "20230001",
                "class_name": "软件工程1班",
                "instructor": "李老师",
                "project_name": "RAG实训项目",
                "report_body": "{{report_body}}"  # 保留占位符
            }
            tpl.render(context_dict)
            
            # 4.2 保存临时渲染结果
            temp_docx = temp_dir / f"temp_rendered_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            tpl.save(temp_docx)
            print(f"\n✅ docxtpl渲染完成: {temp_docx}")
            
            # 4.3 用python-docx加载并处理结构化内容
            doc = Document(temp_docx)
            
            # 4.4 检查渲染后的样式
            print("\n📋 渲染后的样式列表:")
            for style in doc.styles:
                try:
                    style_type = "段落" if style.type == WD_STYLE_TYPE.PARAGRAPH else "字符"
                    print(f"  - {style.name} ({style_type})")
                except:
                    print(f"  - {style.name} (未知类型)")
            
            # 4.5 查找占位符位置
            target_paragraph = None
            for paragraph in doc.paragraphs:
                if "{{report_body}}" in paragraph.text:
                    target_paragraph = paragraph
                    break
            
            if not target_paragraph:
                print("⚠️  未找到占位符，在文档末尾添加内容")
                doc.add_paragraph("报告正文：")
                target_paragraph = doc.paragraphs[-1]
            else:
                print("✅ 找到占位符位置")
            
            # 4.6 测试粗体文本处理
            test_text = "这是一个**重要**的测试文本，包含**粗体**内容。"
            print(f"\n🧪 测试文本: {test_text}")
            
            # 获取可用的样式
            bold_style = get_style(doc, ['bold', 'Bold', 'Strong', 'Normal'])
            normal_style = get_style(doc, ['Normal', '正文'])
            
            print(f"📋 可用样式:")
            print(f"  - 粗体样式: {bold_style}")
            print(f"  - 正文样式: {normal_style}")
            
            # 处理粗体文本
            if '**' in test_text:
                parts = test_text.split('**')
                p = doc.add_paragraph()
                p.style = normal_style or 'Normal'
                
                for i, part in enumerate(parts):
                    if part.strip():
                        if i % 2 == 1:  # 粗体文本
                            run = p.add_run(part)
                            if bold_style and bold_style != 'Normal':
                                try:
                                    if doc.styles[bold_style].type == WD_STYLE_TYPE.CHARACTER:
                                        run.style = bold_style
                                        print(f"✅ 使用字符样式应用粗体: {part}")
                                    else:
                                        run.bold = True
                                        print(f"✅ 使用字体属性应用粗体: {part}")
                                except Exception as e:
                                    run.bold = True
                                    print(f"⚠️  样式应用失败，使用字体属性: {e}")
                            else:
                                run.bold = True
                                print(f"✅ 使用字体属性应用粗体: {part}")
                        else:  # 普通文本
                            p.add_run(part)
            
            # 4.7 保存最终文档
            final_docx = temp_dir / f"final_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            doc.save(final_docx)
            print(f"\n✅ 最终文档保存: {final_docx}")
            
            print("\n🎉 调试完成！请检查生成的文档中粗体文本的显示效果。")
            
        except Exception as e:
            print(f"❌ 调试失败: {e}")
            import traceback
            traceback.print_exc()

if __name__ == "__main__":
    debug_bold_style() 