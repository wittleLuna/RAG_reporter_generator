#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
测试新的富文本渲染方式（docxtpl + python-docx混合渲染）
"""

import os
import tempfile
from pathlib import Path
from docx import Document
from docxtpl import DocxTemplate
from datetime import datetime
from docx.enum.style import WD_STYLE_TYPE

def get_style(doc, style_names):
    """获取存在的样式名，支持多个备选样式名"""
    for style_name in style_names:
        try:
            doc.styles[style_name]
            return style_name
        except KeyError:
            continue
    return None

def create_test_template():
    """创建测试模板"""
    doc = Document()
    
    # 添加标题
    title = doc.add_heading('实训报告', 0)
    
    # 添加基本信息
    doc.add_paragraph('姓名：{{name}}')
    doc.add_paragraph('学号：{{student_id}}')
    doc.add_paragraph('班级：{{class_name}}')
    doc.add_paragraph('指导教师：{{instructor}}')
    doc.add_paragraph('项目名称：{{project_name}}')
    
    # 添加分隔线
    doc.add_paragraph('=' * 50)
    
    # 添加报告正文占位符
    doc.add_paragraph('{{report_body}}')
    
    return doc

def test_rich_text_rendering():
    """测试富文本渲染功能"""
    print("🧪 开始测试富文本渲染功能...")
    
    # 创建临时目录
    with tempfile.TemporaryDirectory() as temp_dir:
        temp_dir = Path(temp_dir)
        
        # 1. 创建测试模板
        template_doc = create_test_template()
        template_path = temp_dir / "test_template.docx"
        template_doc.save(template_path)
        print(f"📄 创建测试模板: {template_path}")
        
        # 2. 创建测试图片目录
        upload_dir = temp_dir / "uploads"
        upload_dir.mkdir()
        
        # 3. 模拟AI生成的markdown内容
        test_markdown = """# 项目概述

这是一个测试项目，用于验证富文本渲染功能。

## 技术栈

- Python 3.8+
- FastAPI
- python-docx
- docxtpl

### 核心组件

项目包含以下**核心组件**：

1. 数据读取模块
2. AI服务模块  
3. 富文本渲染模块

## 代码示例

```python
def hello_world():
    print("Hello, World!")
    return "Success"
```

### 使用方法

这是一个**重要**的使用说明，请仔细阅读。

## 项目结构

项目包含以下主要组件：

1. 数据读取模块
2. AI服务模块  
3. 富文本渲染模块

### 详细说明

每个模块都有其**特定功能**：

- 数据读取：处理文件上传和解析
- AI服务：调用千问API生成内容
- 富文本渲染：将markdown转换为Word格式

## 测试图片

{{image:test_image}}

## 总结

通过本次测试，验证了富文本渲染功能的**正确性**和**完整性**。
"""
        
        print(f"📝 测试markdown内容长度: {len(test_markdown)}字符")
        
        # 4. 模拟渲染过程
        try:
            # 4.1 先用docxtpl渲染模板字段
            tpl = DocxTemplate(template_path)
            context_dict = {
                "name": "张三",
                "student_id": "20230001",
                "class_name": "软件工程1班",
                "instructor": "李老师",
                "project_name": "RAG实训项目",
                "report_body": "{{report_body}}"  # 保留占位符
            }
            tpl.render(context_dict)
            
            # 4.2 保存临时渲染结果
            temp_docx = temp_dir / f"temp_rendered_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            tpl.save(temp_docx)
            print(f"✅ docxtpl渲染完成: {temp_docx}")
            
            # 4.3 用python-docx加载并处理结构化内容
            doc = Document(temp_docx)
            
            # 4.4 查找占位符位置
            target_paragraph = None
            for paragraph in doc.paragraphs:
                if "{{report_body}}" in paragraph.text:
                    target_paragraph = paragraph
                    break
            
            if not target_paragraph:
                print("⚠️  未找到占位符，在文档末尾添加内容")
                doc.add_paragraph("报告正文：")
                target_paragraph = doc.paragraphs[-1]
            else:
                print("✅ 找到占位符位置")
            
            # 4.5 在占位符位置插入结构化内容
            add_structured_content_test(doc, target_paragraph, test_markdown)
            
            # 4.6 保存最终文档
            final_docx = temp_dir / f"final_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            doc.save(final_docx)
            print(f"✅ 最终文档保存: {final_docx}")
            
            # 4.7 验证结果
            verify_rendering_result(final_docx)
            
            print("🎉 富文本渲染测试完成！")
            
        except Exception as e:
            print(f"❌ 测试失败: {e}")
            import traceback
            traceback.print_exc()

def add_structured_content_test(doc, target_paragraph, markdown_text):
    """测试版本的结构化内容添加函数"""
    lines = markdown_text.split('\n')
    in_code_block = False
    code_lines = []
    
    # 获取可用的样式
    code_style = get_style(doc, ['Code', 'Normal'])
    heading1_style = get_style(doc, ['Heading 1', '标题 1', 'Normal'])
    heading2_style = get_style(doc, ['Heading 2', '标题 2', 'Normal'])
    heading3_style = get_style(doc, ['Heading 3', '标题 3', 'Normal'])
    list_style = get_style(doc, ['项目符号', 'Normal'])
    normal_style = get_style(doc, ['Normal', '正文'])
    bold_style = get_style(doc, ['bold', 'Bold', 'Strong', 'Normal'])
    
    print(f"📋 可用样式:")
    print(f"  - 代码样式: {code_style}")
    print(f"  - 一级标题: {heading1_style}")
    print(f"  - 二级标题: {heading2_style}")
    print(f"  - 三级标题: {heading3_style}")
    print(f"  - 列表样式: {list_style}")
    print(f"  - 正文样式: {normal_style}")
    print(f"  - 粗体样式: {bold_style}")
    
    # 删除原占位符段落
    if target_paragraph:
        p = target_paragraph._element
        p.getparent().remove(p)
    
    for line in lines:
        # 代码块识别（markdown风格）
        if line.strip().startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_lines = []
            else:
                # 代码块结束，渲染代码
                in_code_block = False
                code_text = '\n'.join(code_lines)
                # 创建代码块段落，优先使用自定义Code样式
                p = doc.add_paragraph()
                if code_style:
                    p.style = code_style
                    print(f"✅ 使用自定义Code样式: {code_style}")
                else:
                    # 如果没有Code样式，使用默认样式并设置字体
                    run = p.add_run(code_text)
                    run.font.name = 'Consolas'  # 等宽字体
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(64, 64, 64)  # 灰色
                    print("⚠️  使用默认代码块样式")
                p.add_run(code_text)
                code_lines = []
            continue
            
        if in_code_block:
            code_lines.append(line)
            continue
            
        # 图片插入（测试时跳过）
        if '{{image:' in line:
            p = doc.add_paragraph(f"[图片占位符: {line.strip()}]")
            continue
            
        # 项目符号
        if line.strip().startswith('- '):
            p = doc.add_paragraph(line.strip()[2:])
            p.style = list_style or 'Normal'
        # 三级标题
        elif line.strip().startswith('### '):
            p = doc.add_paragraph(line.strip().replace('### ', ''))
            p.style = heading3_style or 'Normal'
            print(f"✅ 添加三级标题: {line.strip().replace('### ', '')}")
        # 二级标题
        elif line.strip().startswith('## '):
            p = doc.add_paragraph(line.strip().replace('## ', ''))
            p.style = heading2_style or 'Normal'
        # 一级标题
        elif line.strip().startswith('# '):
            p = doc.add_paragraph(line.strip().replace('# ', ''))
            p.style = heading1_style or 'Normal'
        # 普通段落（包含粗体处理）
        elif line.strip():
            # 处理粗体文本
            if '**' in line:
                # 分割文本，处理粗体部分
                parts = line.split('**')
                p = doc.add_paragraph()
                p.style = normal_style or 'Normal'
                
                for i, part in enumerate(parts):
                    if part.strip():  # 跳过空字符串
                        if i % 2 == 1:  # 奇数索引是粗体文本
                            run = p.add_run(part)
                            # 优先使用字符样式，如果没有则使用字体属性
                            if bold_style and bold_style != 'Normal':
                                try:
                                    # 检查是否为字符样式
                                    if doc.styles[bold_style].type == WD_STYLE_TYPE.CHARACTER:
                                        run.style = bold_style
                                        print(f"✅ 使用字符样式应用粗体: {part}")
                                    else:
                                        # 如果不是字符样式，使用字体属性
                                        run.bold = True
                                        print(f"✅ 使用字体属性应用粗体: {part}")
                                except Exception as e:
                                    # 如果样式应用失败，使用字体属性
                                    run.bold = True
                                    print(f"⚠️  样式应用失败，使用字体属性: {e}")
                            else:
                                # 没有找到粗体样式，使用字体属性
                                run.bold = True
                                print(f"✅ 使用字体属性应用粗体: {part}")
                        else:  # 偶数索引是普通文本
                            p.add_run(part)
            else:
                p = doc.add_paragraph(line.strip())
                p.style = normal_style or 'Normal'

def verify_rendering_result(doc_path):
    """验证渲染结果"""
    print("\n🔍 验证渲染结果...")
    
    doc = Document(doc_path)
    
    # 检查段落数量
    paragraph_count = len(doc.paragraphs)
    print(f"  - 总段落数: {paragraph_count}")
    
    # 检查标题
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith('标题')]
    print(f"  - 标题数量: {len(headings)}")
    for heading in headings:
        print(f"    - {heading}")
    
    # 检查项目符号
    bullets = [p.text for p in doc.paragraphs if p.style.name == '项目符号']
    print(f"  - 项目符号数量: {len(bullets)}")
    
    # 检查代码块
    code_blocks = [p.text for p in doc.paragraphs if 'def hello_world' in p.text]
    print(f"  - 代码块数量: {len(code_blocks)}")
    
    print("✅ 渲染结果验证完成")

# 导入必要的模块
from docx.shared import Pt, RGBColor

if __name__ == "__main__":
    test_rich_text_rendering() 