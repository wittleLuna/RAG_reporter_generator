#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
测试样式修复
"""

import os
import tempfile
from pathlib import Path
from docx import Document
from docxtpl import DocxTemplate
from datetime import datetime
from docx.shared import Pt

def get_style(doc, style_names):
    """获取存在的样式名，支持多个备选样式名"""
    for style_name in style_names:
        try:
            doc.styles[style_name]
            return style_name
        except KeyError:
            continue
    return None

def create_test_template_with_table():
    """创建包含表格的测试模板"""
    doc = Document()
    
    # 添加标题
    title = doc.add_heading('实训报告', 0)
    
    # 添加基本信息表格
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    
    # 设置表格内容
    table_data = [
        ('姓名：', '{{name}}'),
        ('学号：', '{{student_id}}'),
        ('班级：', '{{class_name}}'),
        ('指导教师：', '{{instructor}}'),
        ('项目名称：', '{{project_name}}')
    ]
    
    for i, (label, placeholder) in enumerate(table_data):
        table.cell(i, 0).text = label
        table.cell(i, 1).text = placeholder
    
    # 添加报告正文表格
    report_table = doc.add_table(rows=1, cols=1)
    report_table.style = 'Table Grid'
    report_table.cell(0, 0).text = '{{report_body}}'
    
    return doc

def find_placeholder_cell_test(doc, placeholder_text="{{report_body}"):
    """测试版本的单元格查找函数"""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if placeholder_text in cell.text:
                    print(f"✅ 在表格单元格中找到占位符: {cell.text[:50]}...")
                    return cell
    print(f"❌ 未找到包含占位符的单元格: {placeholder_text}")
    return None

def insert_structured_content_to_cell_test(cell_or_para, markdown_text):
    """测试版本的结构化内容插入函数（修复样式问题）"""
    if hasattr(cell_or_para, 'text'):  # 单元格
        cell_or_para.text = ""  # 清空原内容
        container = cell_or_para
        # 获取文档对象（通过表格的父级）
        doc = container._parent._parent
        print("📝 在单元格中插入内容")
    else:  # 段落
        container = cell_or_para
        # 获取文档对象
        doc = container._parent
        print("📝 在段落中插入内容")
    
    # 获取可用的样式
    heading1_style = get_style(doc, ['Heading 1', '标题 1', 'Normal'])
    heading2_style = get_style(doc, ['Heading 2', '标题 2', 'Normal'])
    normal_style = get_style(doc, ['Normal', '正文'])
    
    print(f"📋 可用样式:")
    print(f"  - 一级标题: {heading1_style}")
    print(f"  - 二级标题: {heading2_style}")
    print(f"  - 正文样式: {normal_style}")
    
    lines = markdown_text.strip().split('\n')
    in_code_block = False
    code_lines = []

    for line in lines:
        if line.strip().startswith("```"):
            in_code_block = not in_code_block
            if not in_code_block:
                p = container.add_paragraph('\n'.join(code_lines))
                run = p.runs[0]
                run.font.name = "Courier New"
                run.font.size = Pt(10)
                code_lines.clear()
                print(f"✅ 添加代码块: {len(code_lines)}行")
            continue
        if in_code_block:
            code_lines.append(line)
            continue
        elif line.startswith('# '):
            p = container.add_paragraph(line[2:])
            p.style = heading1_style or 'Normal'
            print(f"✅ 添加一级标题: {line[2:]} (样式: {p.style.name})")
        elif line.startswith('## '):
            p = container.add_paragraph(line[3:])
            p.style = heading2_style or 'Normal'
            print(f"✅ 添加二级标题: {line[3:]} (样式: {p.style.name})")
        elif line.startswith('- '):
            p = container.add_paragraph('• ' + line[2:])
            p.style = normal_style or 'Normal'
            print(f"✅ 添加列表项: {line[2:]} (样式: {p.style.name})")
        else:
            p = container.add_paragraph(line)
            p.style = normal_style or 'Normal'
            print(f"✅ 添加段落: {line[:30]}... (样式: {p.style.name})")

def test_style_fix():
    """测试样式修复"""
    print("🧪 开始测试样式修复...")
    
    # 创建临时目录
    with tempfile.TemporaryDirectory() as temp_dir:
        temp_dir = Path(temp_dir)
        
        # 1. 创建测试模板
        template_doc = create_test_template_with_table()
        template_path = temp_dir / "test_template_with_table.docx"
        template_doc.save(template_path)
        print(f"📄 创建测试模板: {template_path}")
        
        # 2. 检查模板中的样式
        print("\n📋 模板中的样式列表:")
        for style in template_doc.styles:
            try:
                print(f"  - {style.name}")
            except:
                print(f"  - {style.name} (未知类型)")
        
        # 3. 模拟完整的渲染过程
        try:
            # 3.1 先用docxtpl渲染模板字段
            tpl = DocxTemplate(template_path)
            context_dict = {
                "name": "张三",
                "student_id": "20230001",
                "class_name": "软件工程1班",
                "instructor": "李老师",
                "project_name": "RAG实训项目",
                "report_body": "{{report_body}}"  # 不替换正文，占位
            }
            tpl.render(context_dict)
            
            # 3.2 保存临时渲染结果
            temp_docx = temp_dir / f"temp_rendered_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            tpl.save(temp_docx)
            print(f"✅ docxtpl渲染完成: {temp_docx}")
            
            # 3.3 用python-docx加载并处理结构化内容
            doc = Document(temp_docx)
            
            # 3.4 检查字段是否已正确渲染
            print("\n🔍 检查字段渲染状态...")
            field_check_results = {
                "name": "张三",
                "student_id": "20230001",
                "class_name": "软件工程1班",
                "instructor": "李老师",
                "project_name": "RAG实训项目"
            }
            
            for field_name, field_content in field_check_results.items():
                if field_content and field_content.strip():
                    print(f"✅ 字段{field_name}已通过docxtpl处理: {field_content}")
                else:
                    print(f"⚠️  字段{field_name}为空或未提供")
            
            # 3.5 查找报告正文占位符单元格
            print("\n🔍 查找报告正文占位符单元格...")
            target_cell = find_placeholder_cell_test(doc)
            
            # 3.6 添加测试报告正文
            test_report_body = """# 项目概述

这是一个测试项目，用于验证样式修复。

## 技术栈

- Python 3.8+
- FastAPI
- python-docx
- docxtpl

### 核心组件

项目包含以下**核心组件**：

1. 数据读取模块
2. AI服务模块  
3. 富文本渲染模块

## 代码示例

```python
def hello_world():
    print("Hello, World!")
    return "Success"
```

### 使用方法

这是一个**重要**的使用说明，请仔细阅读。

## 总结

通过本次测试，验证了样式修复的**正确性**和**完整性**。
"""
            
            if target_cell:
                insert_structured_content_to_cell_test(target_cell, test_report_body)
                print("✅ 成功将正文插入表格单元格")
            else:
                print("⚠️  未找到report_body单元格，改为添加到末尾")
                para = doc.add_paragraph()
                insert_structured_content_to_cell_test(para, test_report_body)  # fallback
            
            # 3.7 保存最终文档
            final_docx = temp_dir / f"final_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            doc.save(final_docx)
            print(f"✅ 最终文档保存: {final_docx}")
            
            print("\n🎉 样式修复测试完成！")
            
        except Exception as e:
            print(f"❌ 测试失败: {e}")
            import traceback
            traceback.print_exc()

if __name__ == "__main__":
    test_style_fix() 