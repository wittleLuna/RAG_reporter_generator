#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
创建示例模板文件
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

def create_cover_template():
    """创建封面模板"""
    doc = Document()
    
    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.5)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
    
    # 添加标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("实训报告")
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    
    # 添加空行
    doc.add_paragraph()
    doc.add_paragraph()
    
    # 添加信息表格
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    
    # 设置表格内容
    info_data = [
        ("姓名", "{{name}}"),
        ("学号", "{{student_id}}"),
        ("班级", "{{class_name}}"),
        ("指导老师", "{{instructor}}"),
        ("项目名称", "{{project_name}}")
    ]
    
    for i, (label, value) in enumerate(info_data):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        
        # 设置标签列样式
        label_cell = row.cells[0]
        label_cell.paragraphs[0].runs[0].font.bold = True
        label_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 设置值列样式
        value_cell = row.cells[1]
        value_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # 调整表格宽度
    table.columns[0].width = Inches(1.5)
    table.columns[1].width = Inches(4)
    
    # 添加空行
    doc.add_paragraph()
    doc.add_paragraph()
    
    # 添加日期
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run("日期：{{date}}")
    date_run.font.size = Pt(12)
    
    # 保存文件
    if not os.path.exists("templates"):
        os.makedirs("templates")
    
    doc.save("templates/cover_template.docx")
    print("✅ 封面模板已创建: templates/cover_template.docx")

def create_body_template():
    """创建正文模板"""
    doc = Document()
    
    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
    
    # 添加标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("{{project_name}}")
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    
    # 添加空行
    doc.add_paragraph()
    
    # 添加扩展信息表格（共10行2列）
    info_table = doc.add_table(rows=10, cols=2)
    info_table.style = 'Table Grid'
    
    info_data = [
        ("姓名", "{{name}}"),
        ("学号", "{{student_id}}"),
        ("班级", "{{class_name}}"),
        ("指导老师", "{{instructor}}"),
        ("教材", "{{textbook}}"),
        ("实训室", "{{lab}}"),
        ("完成日期", "{{finish_date}}"),
        ("设计要求", "{{design_requirements}}"),
        ("所用知识与技术", "{{knowledge_and_tech}}"),
        ("完成情况", "{{completion}}"),
        ("自我说明", "{{self_statement}}"),
    ]
    # 由于表格只创建了10行，info_data有11项，需修正为11行
    if len(info_table.rows) < len(info_data):
        for _ in range(len(info_data) - len(info_table.rows)):
            info_table.add_row()
    
    for i, (label, value) in enumerate(info_data):
        row = info_table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        # 设置样式
        label_cell = row.cells[0]
        label_cell.paragraphs[0].runs[0].font.bold = True
        label_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        value_cell = row.cells[1]
        value_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # 调整表格宽度
    info_table.columns[0].width = Inches(1.5)
    info_table.columns[1].width = Inches(4)
    
    # 添加空行
    doc.add_paragraph()
    doc.add_paragraph()
    
    # 添加正文占位符
    body_para = doc.add_paragraph()
    body_para.add_run("{{report_body}}")
    
    # 保存文件
    if not os.path.exists("templates"):
        os.makedirs("templates")
    
    doc.save("templates/body_template.docx")
    print("✅ 正文模板已创建: templates/body_template.docx")

def create_combined_template():
    """创建组合模板（封面+正文）"""
    doc = Document()
    
    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
    
    # 添加标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("实训报告")
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    
    # 添加空行
    doc.add_paragraph()
    doc.add_paragraph()
    
    # 添加信息表格
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    
    # 设置表格内容
    info_data = [
        ("姓名", "{{name}}"),
        ("学号", "{{student_id}}"),
        ("班级", "{{class_name}}"),
        ("指导老师", "{{instructor}}"),
        ("项目名称", "{{project_name}}")
    ]
    
    for i, (label, value) in enumerate(info_data):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        
        # 设置标签列样式
        label_cell = row.cells[0]
        label_cell.paragraphs[0].runs[0].font.bold = True
        label_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 设置值列样式
        value_cell = row.cells[1]
        value_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # 调整表格宽度
    table.columns[0].width = Inches(1.5)
    table.columns[1].width = Inches(4)
    
    # 添加空行
    doc.add_paragraph()
    doc.add_paragraph()
    
    # 添加正文占位符
    body_para = doc.add_paragraph()
    body_para.add_run("{{report_body}}")
    
    # 保存文件
    if not os.path.exists("templates"):
        os.makedirs("templates")
    
    doc.save("templates/combined_template.docx")
    print("✅ 组合模板已创建: templates/combined_template.docx")

if __name__ == "__main__":
    print("🚀 开始创建示例模板文件...")
    
    create_cover_template()
    create_body_template()
    create_combined_template()
    
    print("\n📝 模板文件说明：")
    print("1. cover_template.docx - 封面模板（仅包含基本信息字段）")
    print("2. body_template.docx - 正文模板（包含报告正文占位符）")
    print("3. combined_template.docx - 组合模板（封面+正文，兼容旧版本）")
    print("\n💡 使用建议：")
    print("- 推荐使用 cover_template.docx + body_template.docx 分离模式")
    print("- 系统会自动识别并合并封面和正文")
    print("- 如果只上传一个模板，系统会将其作为正文模板使用") 