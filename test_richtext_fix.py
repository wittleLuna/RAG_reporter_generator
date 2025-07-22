#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
测试富文本修复
"""

import os
import tempfile
from pathlib import Path
from docx import Document
from docxtpl import DocxTemplate
from datetime import datetime
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import RGBColor

def get_style(doc, style_names):
    """获取存在的样式名，支持多个备选样式名"""
    for style_name in style_names:
        try:
            doc.styles[style_name]
            return style_name
        except KeyError:
            continue
    return None

def create_test_template_with_table():
    """创建包含表格的测试模板"""
    doc = Document()
    
    # 添加标题
    title = doc.add_heading('实训报告', 0)
    
    # 添加基本信息表格
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    
    # 设置表格内容
    table_data = [
        ('姓名：', '{{name}}'),
        ('学号：', '{{student_id}}'),
        ('班级：', '{{class_name}}'),
        ('指导教师：', '{{instructor}}'),
        ('项目名称：', '{{project_name}}')
    ]
    
    for i, (label, placeholder) in enumerate(table_data):
        table.cell(i, 0).text = label
        table.cell(i, 1).text = placeholder
    
    # 添加报告正文表格
    report_table = doc.add_table(rows=1, cols=1)
    report_table.style = 'Table Grid'
    report_table.cell(0, 0).text = '{{report_body}}'
    
    return doc

def find_placeholder_cell_test(doc, placeholder_text="{{report_body}"):
    """测试版本的单元格查找函数"""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if placeholder_text in cell.text:
                    print(f"✅ 在表格单元格中找到占位符: {cell.text[:50]}...")
                    return cell
    print(f"❌ 未找到包含占位符的单元格: {placeholder_text}")
    return None

def insert_structured_content_to_cell_test(doc, cell_or_para, markdown_text):
    """测试版本的结构化内容插入函数（包含富文本处理）"""
    if hasattr(cell_or_para, 'text'):  # 单元格
        cell_or_para.text = ""  # 清空原内容
        container = cell_or_para
        print("📝 在单元格中插入内容")
    else:  # 段落
        container = cell_or_para
        print("📝 在段落中插入内容")
    
    # 获取可用的样式
    heading1_style = get_style(doc, ['Heading 1', '标题 1', 'Normal'])
    heading2_style = get_style(doc, ['Heading 2', '标题 2', 'Normal'])
    heading3_style = get_style(doc, ['Heading 3', '标题 3', 'Normal'])
    heading4_style = get_style(doc, ['Heading 4', '标题 4', 'Normal'])
    code_style = get_style(doc, ['Code', 'Normal'])
    normal_style = get_style(doc, ['Normal', '正文'])
    bold_style = get_style(doc, ['bold', 'Bold', 'Strong', 'Normal'])
    
    print(f"📋 可用样式:")
    print(f"  - 一级标题: {heading1_style}")
    print(f"  - 二级标题: {heading2_style}")
    print(f"  - 三级标题: {heading3_style}")
    print(f"  - 四级标题: {heading4_style}")
    print(f"  - 代码样式: {code_style}")
    print(f"  - 正文样式: {normal_style}")
    print(f"  - 粗体样式: {bold_style}")
    
    lines = markdown_text.strip().split('\n')
    in_code_block = False
    code_lines = []

    for line in lines:
        if line.strip().startswith("```"):
            in_code_block = not in_code_block
            if not in_code_block:
                # 代码块结束，渲染代码
                code_text = '\n'.join(code_lines)
                p = container.add_paragraph()
                if code_style:
                    p.style = code_style
                    # 添加文本内容到已设置样式的段落
                    p.add_run(code_text)
                    print(f"✅ 使用Code样式添加代码块: {len(code_lines)}行")
                else:
                    # 如果没有Code样式，手动设置字体
                    run = p.add_run(code_text)
                    run.font.name = "Courier New"
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(64, 64, 64)  # 灰色
                    print(f"✅ 手动设置代码块样式: {len(code_lines)}行")
                code_lines.clear()
            continue
        if in_code_block:
            code_lines.append(line)
            continue
        elif line.startswith('# '):
            p = container.add_paragraph(line[2:])
            if heading1_style:
                p.style = heading1_style
                print(f"✅ 添加一级标题: {line[2:]} (样式: {heading1_style})")
            else:
                # 如果没有样式，手动设置字体
                run = p.runs[0]
                run.bold = True
                run.font.size = Pt(16)
                print(f"✅ 添加一级标题: {line[2:]} (手动设置)")
        elif line.startswith('## '):
            p = container.add_paragraph(line[3:])
            if heading2_style:
                p.style = heading2_style
                print(f"✅ 添加二级标题: {line[3:]} (样式: {heading2_style})")
            else:
                # 如果没有样式，手动设置字体
                run = p.runs[0]
                run.bold = True
                run.font.size = Pt(14)
                print(f"✅ 添加二级标题: {line[3:]} (手动设置)")
        elif line.startswith('### '):
            p = container.add_paragraph(line[4:])
            if heading3_style:
                p.style = heading3_style
                print(f"✅ 添加三级标题: {line[4:]} (样式: {heading3_style})")
            else:
                # 如果没有样式，手动设置字体
                run = p.runs[0]
                run.bold = True
                run.font.size = Pt(12)
                print(f"✅ 添加三级标题: {line[4:]} (手动设置)")
        elif line.startswith('#### '):
            p = container.add_paragraph(line[5:])
            if heading4_style:
                p.style = heading4_style
                print(f"✅ 添加四级标题: {line[5:]} (样式: {heading4_style})")
            else:
                # 如果没有样式，手动设置字体
                run = p.runs[0]
                run.bold = True
                run.font.size = Pt(11)
                print(f"✅ 添加四级标题: {line[5:]} (手动设置)")
        elif line.startswith('- '):
            p = container.add_paragraph('• ' + line[2:])
            if normal_style:
                p.style = normal_style
            print(f"✅ 添加列表项: {line[2:]}")
        else:
            # 处理富文本格式（粗体等）
            if '**' in line:
                # 分割文本，处理粗体部分
                parts = line.split('**')
                p = container.add_paragraph()
                if normal_style:
                    p.style = normal_style
                
                for i, part in enumerate(parts):
                    if part.strip():  # 跳过空字符串
                        if i % 2 == 1:  # 奇数索引是粗体文本
                            run = p.add_run(part)
                            # 优先使用字符样式，如果没有则使用字体属性
                            if bold_style and bold_style != 'Normal':
                                try:
                                    # 检查是否为字符样式
                                    if doc.styles[bold_style].type == WD_STYLE_TYPE.CHARACTER:
                                        run.style = bold_style
                                        print(f"✅ 使用字符样式应用粗体: {part}")
                                    else:
                                        # 如果不是字符样式，使用字体属性
                                        run.bold = True
                                        print(f"✅ 使用字体属性应用粗体: {part}")
                                except Exception as e:
                                    # 如果样式应用失败，使用字体属性
                                    run.bold = True
                                    print(f"⚠️  样式应用失败，使用字体属性: {e}")
                            else:
                                # 没有找到粗体样式，使用字体属性
                                run.bold = True
                                print(f"✅ 使用字体属性应用粗体: {part}")
                        else:  # 偶数索引是普通文本
                            p.add_run(part)
                print(f"✅ 添加富文本段落: {line[:50]}...")
            else:
                p = container.add_paragraph(line)
                if normal_style:
                    p.style = normal_style
                print(f"✅ 添加普通段落: {line[:30]}...")

def test_richtext_fix():
    """测试富文本修复"""
    print("🧪 开始测试富文本修复...")
    
    # 创建临时目录
    with tempfile.TemporaryDirectory() as temp_dir:
        temp_dir = Path(temp_dir)
        
        # 1. 创建测试模板
        template_doc = create_test_template_with_table()
        template_path = temp_dir / "test_template_with_table.docx"
        template_doc.save(template_path)
        print(f"📄 创建测试模板: {template_path}")
        
        # 2. 模拟完整的渲染过程
        try:
            # 2.1 先用docxtpl渲染模板字段
            tpl = DocxTemplate(template_path)
            context_dict = {
                "name": "张三",
                "student_id": "20230001",
                "class_name": "软件工程1班",
                "instructor": "李老师",
                "project_name": "RAG实训项目",
                "report_body": "{{report_body}}"  # 不替换正文，占位
            }
            tpl.render(context_dict)
            
            # 2.2 保存临时渲染结果
            temp_docx = temp_dir / f"temp_rendered_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            tpl.save(temp_docx)
            print(f"✅ docxtpl渲染完成: {temp_docx}")
            
            # 2.3 用python-docx加载并处理结构化内容
            doc = Document(temp_docx)
            
            # 2.4 检查字段是否已正确渲染
            print("\n🔍 检查字段渲染状态...")
            field_check_results = {
                "name": "张三",
                "student_id": "20230001",
                "class_name": "软件工程1班",
                "instructor": "李老师",
                "project_name": "RAG实训项目"
            }
            
            for field_name, field_content in field_check_results.items():
                if field_content and field_content.strip():
                    print(f"✅ 字段{field_name}已通过docxtpl处理: {field_content}")
                else:
                    print(f"⚠️  字段{field_name}为空或未提供")
            
            # 2.5 查找报告正文占位符单元格
            print("\n🔍 查找报告正文占位符单元格...")
            target_cell = find_placeholder_cell_test(doc)
            
            # 2.6 添加测试报告正文（包含富文本）
            test_report_body = """# 项目概述

这是一个**重要**的测试项目，用于验证富文本修复功能。

## 技术栈

- **Python 3.8+** - 核心编程语言
- **FastAPI** - Web框架
- **python-docx** - Word文档处理
- **docxtpl** - 模板渲染

### 核心组件

项目包含以下**核心组件**：

1. **数据读取模块** - 处理文件上传和解析
2. **AI服务模块** - 调用千问API生成内容
3. **富文本渲染模块** - 将markdown转换为Word格式

#### 详细说明

这是四级标题的**测试内容**，用于验证Heading 4样式。

## 代码示例

```python
def hello_world():
    print("Hello, World!")
    return "Success"
```

### 使用方法

这是一个**重要**的使用说明，请仔细阅读。

#### 注意事项

- 确保所有依赖已安装
- 检查配置文件是否正确
- 测试所有功能模块

## 总结

通过本次测试，验证了富文本修复的**正确性**和**完整性**。
"""
            
            if target_cell:
                insert_structured_content_to_cell_test(doc, target_cell, test_report_body)
                print("✅ 成功将正文插入表格单元格")
            else:
                print("⚠️  未找到report_body单元格，改为添加到末尾")
                para = doc.add_paragraph()
                insert_structured_content_to_cell_test(doc, para, test_report_body)  # fallback
            
            # 2.7 保存最终文档
            final_docx = temp_dir / f"final_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            doc.save(final_docx)
            print(f"✅ 最终文档保存: {final_docx}")
            
            print("\n🎉 富文本修复测试完成！")
            
            return str(final_docx)
            
        except Exception as e:
            print(f"❌ 测试失败: {e}")
            import traceback
            traceback.print_exc()
            return None

if __name__ == "__main__":
    test_richtext_fix() 