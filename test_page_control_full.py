#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
完整测试页面控制功能 - 包括实际文档生成
"""

import os
import time
from app import build_prompt, generate_report
from docx import Document
from docx.shared import Inches

def count_pages(docx_path):
    """估算文档页数"""
    try:
        doc = Document(docx_path)
        
        # 简单估算：每页约500-800字符
        total_chars = 0
        for paragraph in doc.paragraphs:
            total_chars += len(paragraph.text)
        
        # 估算页数（每页约600字符）
        estimated_pages = max(1, total_chars // 600)
        
        # 检查是否有分页符
        page_breaks = 0
        for paragraph in doc.paragraphs:
            if paragraph.runs:
                for run in paragraph.runs:
                    if hasattr(run, '_element') and run._element.xml.find('w:br') != -1:
                        page_breaks += 1
        
        return max(estimated_pages, page_breaks + 1)
        
    except Exception as e:
        print(f"  ⚠️  页数估算失败: {e}")
        return 0

def test_page_control_generation():
    """测试页面控制文档生成"""
    print("🧪 测试页面控制文档生成...")
    
    # 测试用例
    test_cases = [
        (2, "短报告"),
        (5, "中等报告"), 
        (10, "长报告")
    ]
    
    query = "请生成一个关于人工智能在医疗领域应用的实训报告"
    context_text = """
    人工智能在医疗领域的应用越来越广泛，包括：
    1. 医学影像诊断：AI可以辅助医生识别X光片、CT扫描等医学影像中的异常
    2. 药物研发：AI可以加速新药的发现和开发过程
    3. 个性化医疗：基于患者数据的个性化治疗方案
    4. 医疗机器人：手术机器人和康复机器人
    5. 智能健康监测：可穿戴设备和远程监测系统
    """
    
    for target_pages, description in test_cases:
        print(f"\n📄 生成 {description} (目标页数: {target_pages})")
        
        try:
            # 生成文档
            start_time = time.time()
            docx_path = generate_report(query, context_text, target_pages=target_pages)
            generation_time = time.time() - start_time
            
            if docx_path and os.path.exists(docx_path):
                # 估算页数
                estimated_pages = count_pages(docx_path)
                
                print(f"  ✅ 文档生成成功")
                print(f"  📁 文件路径: {docx_path}")
                print(f"  📊 估算页数: {estimated_pages}")
                print(f"  ⏱️  生成时间: {generation_time:.2f}秒")
                
                # 检查页数是否符合预期
                if target_pages <= 5:
                    if estimated_pages >= target_pages * 0.7 and estimated_pages <= target_pages * 1.5:
                        print(f"  ✅ 页数符合预期 (±50%)")
                    else:
                        print(f"  ⚠️  页数可能不符合预期")
                else:
                    if estimated_pages >= target_pages * 0.5:
                        print(f"  ✅ 页数基本符合预期")
                    else:
                        print(f"  ⚠️  页数可能偏少")
                        
            else:
                print(f"  ❌ 文档生成失败")
                
        except Exception as e:
            print(f"  ❌ 生成失败: {e}")

def test_prompt_analysis():
    """分析提示词内容"""
    print("\n🔍 分析提示词内容...")
    
    query = "测试查询"
    context_text = "测试上下文"
    
    # 测试不同页数的提示词
    for pages in [2, 5, 10, 20]:
        prompt = build_prompt(query, context_text, target_pages=pages)
        
        print(f"\n📝 {pages}页报告提示词分析:")
        print(f"  总长度: {len(prompt)} 字符")
        
        # 检查关键要素 - 修正字符串匹配
        checks = [
            ("页面控制要求", "页面控制要求" in prompt),
            ("目标页数", f"对应 {pages} 页" in prompt),
            ("内容结构", "报告结构" in prompt or "章节" in prompt),
            ("详细程度", "详细" in prompt or "深入" in prompt),
            ("格式要求", "格式" in prompt or "排版" in prompt)
        ]
        
        for check_name, result in checks:
            status = "✅" if result else "❌"
            print(f"  {status} {check_name}")
        
        # 显示页面控制部分
        if "页面控制要求" in prompt:
            start_idx = prompt.find("页面控制要求")
            end_idx = prompt.find("查询问题：")
            if end_idx == -1:
                end_idx = prompt.find("相关资料")
            if end_idx == -1:
                end_idx = start_idx + 500
            
            page_control_text = prompt[start_idx:end_idx].strip()
            print(f"  📋 页面控制内容: {page_control_text[:100]}...")

def main():
    """主测试函数"""
    print("🚀 开始完整页面控制功能测试")
    print("=" * 60)
    
    try:
        # 1. 分析提示词
        test_prompt_analysis()
        
        # 2. 测试文档生成
        test_page_control_generation()
        
        print("\n" + "=" * 60)
        print("✅ 完整页面控制功能测试完成")
        
    except Exception as e:
        print(f"\n❌ 测试失败: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main() 