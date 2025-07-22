#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
专门测试四级标题
"""

import tempfile
from pathlib import Path
from docx import Document
from docxtpl import DocxTemplate
from datetime import datetime
from docx.shared import Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE

def get_style(doc, style_names):
    """获取存在的样式名，支持多个备选样式名"""
    for style_name in style_names:
        try:
            doc.styles[style_name]
            return style_name
        except KeyError:
            continue
    return None

def create_test_template():
    """创建测试模板"""
    doc = Document()
    
    # 添加标题
    title = doc.add_heading('四级标题测试', 0)
    
    # 添加报告正文表格
    report_table = doc.add_table(rows=1, cols=1)
    report_table.style = 'Table Grid'
    report_table.cell(0, 0).text = '{{report_body}}'
    
    return doc

def insert_structured_content_to_cell_test(doc, cell_or_para, markdown_text):
    """测试版本的结构化内容插入函数"""
    if hasattr(cell_or_para, 'text'):  # 单元格
        cell_or_para.text = ""  # 清空原内容
        container = cell_or_para
        print("📝 在单元格中插入内容")
    else:  # 段落
        container = cell_or_para
        print("📝 在段落中插入内容")
    
    # 获取可用的样式
    heading1_style = get_style(doc, ['Heading 1', '标题 1', 'Normal'])
    heading2_style = get_style(doc, ['Heading 2', '标题 2', 'Normal'])
    heading3_style = get_style(doc, ['Heading 3', '标题 3', 'Normal'])
    heading4_style = get_style(doc, ['Heading 4', '标题 4', 'Normal'])
    normal_style = get_style(doc, ['Normal', '正文'])
    bold_style = get_style(doc, ['bold', 'Bold', 'Strong', 'Normal'])
    
    print(f"📋 可用样式:")
    print(f"  - 一级标题: {heading1_style}")
    print(f"  - 二级标题: {heading2_style}")
    print(f"  - 三级标题: {heading3_style}")
    print(f"  - 四级标题: {heading4_style}")
    print(f"  - 正文样式: {normal_style}")
    print(f"  - 粗体样式: {bold_style}")
    
    lines = markdown_text.strip().split('\n')

    for line in lines:
        print(f"🔍 处理行: {line[:30]}...")
        
        if line.startswith('# '):
            p = container.add_paragraph(line[2:])
            if heading1_style:
                p.style = heading1_style
                print(f"✅ 添加一级标题: {line[2:]} (样式: {heading1_style})")
            else:
                run = p.runs[0]
                run.bold = True
                run.font.size = Pt(16)
                print(f"✅ 添加一级标题: {line[2:]} (手动设置)")
        elif line.startswith('## '):
            p = container.add_paragraph(line[3:])
            if heading2_style:
                p.style = heading2_style
                print(f"✅ 添加二级标题: {line[3:]} (样式: {heading2_style})")
            else:
                run = p.runs[0]
                run.bold = True
                run.font.size = Pt(14)
                print(f"✅ 添加二级标题: {line[3:]} (手动设置)")
        elif line.startswith('### '):
            p = container.add_paragraph(line[4:])
            if heading3_style:
                p.style = heading3_style
                print(f"✅ 添加三级标题: {line[4:]} (样式: {heading3_style})")
            else:
                run = p.runs[0]
                run.bold = True
                run.font.size = Pt(12)
                print(f"✅ 添加三级标题: {line[4:]} (手动设置)")
        elif line.startswith('#### '):
            p = container.add_paragraph(line[5:])
            if heading4_style:
                p.style = heading4_style
                print(f"✅ 添加四级标题: {line[5:]} (样式: {heading4_style})")
            else:
                run = p.runs[0]
                run.bold = True
                run.font.size = Pt(11)
                print(f"✅ 添加四级标题: {line[5:]} (手动设置)")
        elif line.startswith('- '):
            p = container.add_paragraph('• ' + line[2:])
            if normal_style:
                p.style = normal_style
            print(f"✅ 添加列表项: {line[2:]}")
        elif line.strip():
            # 处理富文本格式（粗体等）
            if '**' in line:
                parts = line.split('**')
                p = container.add_paragraph()
                if normal_style:
                    p.style = normal_style
                
                for i, part in enumerate(parts):
                    if part.strip():
                        if i % 2 == 1:  # 奇数索引是粗体文本
                            run = p.add_run(part)
                            if bold_style and bold_style != 'Normal':
                                try:
                                    if doc.styles[bold_style].type == WD_STYLE_TYPE.CHARACTER:
                                        run.style = bold_style
                                        print(f"✅ 使用字符样式应用粗体: {part}")
                                    else:
                                        run.bold = True
                                        print(f"✅ 使用字体属性应用粗体: {part}")
                                except Exception as e:
                                    run.bold = True
                                    print(f"⚠️  样式应用失败，使用字体属性: {e}")
                            else:
                                run.bold = True
                                print(f"✅ 使用字体属性应用粗体: {part}")
                        else:
                            p.add_run(part)
                print(f"✅ 添加富文本段落: {line[:50]}...")
            else:
                p = container.add_paragraph(line)
                if normal_style:
                    p.style = normal_style
                print(f"✅ 添加普通段落: {line[:30]}...")

def test_heading4():
    """测试四级标题"""
    print("🧪 开始测试四级标题...")
    
    with tempfile.TemporaryDirectory() as temp_dir:
        temp_dir = Path(temp_dir)
        
        # 1. 创建测试模板
        template_doc = create_test_template()
        template_path = temp_dir / "test_template.docx"
        template_doc.save(template_path)
        print(f"📄 创建测试模板: {template_path}")
        
        try:
            # 2. 先用docxtpl渲染模板字段
            tpl = DocxTemplate(template_path)
            context_dict = {
                "report_body": "{{report_body}}"
            }
            tpl.render(context_dict)
            
            # 3. 保存临时渲染结果
            temp_docx = temp_dir / f"temp_rendered_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            tpl.save(temp_docx)
            print(f"✅ docxtpl渲染完成: {temp_docx}")
            
            # 4. 用python-docx加载并处理结构化内容
            doc = Document(temp_docx)
            
            # 5. 查找报告正文占位符单元格
            target_cell = None
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if "{{report_body}}" in cell.text:
                            target_cell = cell
                            print("✅ 在表格单元格中找到占位符")
                            break
                    if target_cell:
                        break
                if target_cell:
                    break
            
            # 6. 添加测试内容（专门测试四级标题）
            test_content = """# 一级标题测试

这是**一级标题**的内容。

## 二级标题测试

这是**二级标题**的内容。

### 三级标题测试

这是**三级标题**的内容。

#### 四级标题测试

这是**四级标题**的内容，应该使用Heading 4样式。

#### 另一个四级标题

这是另一个**四级标题**的内容。

## 总结

- 一级标题：#
- 二级标题：##
- 三级标题：###
- 四级标题：####
"""
            
            if target_cell:
                insert_structured_content_to_cell_test(doc, target_cell, test_content)
                print("✅ 成功将内容插入表格单元格")
            else:
                print("⚠️  未找到report_body单元格")
            
            # 7. 保存最终文档
            final_docx = temp_dir / f"heading4_test_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            doc.save(final_docx)
            print(f"✅ 最终文档保存: {final_docx}")
            
            print("\n🎉 四级标题测试完成！")
            return str(final_docx)
            
        except Exception as e:
            print(f"❌ 测试失败: {e}")
            import traceback
            traceback.print_exc()
            return None

if __name__ == "__main__":
    test_heading4() 